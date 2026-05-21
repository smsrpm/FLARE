"""
flare_ui.py   -   Streamlit web interface for flare_sim.py
=====================================================================
Run with:
    streamlit run flare_ui.py

Requirements:
    pip install streamlit plotly openpyxl pandas numpy
    flare_sim.py and *_in.xlsx files in same directory.
"""

import io
import math
import os
import re
import sys
import json
import shutil
import subprocess
import time
import threading
import queue
from datetime import datetime
from pathlib import Path

import numpy as np
import pandas as pd
import plotly.graph_objects as go
import plotly.express as px
from plotly.subplots import make_subplots
import requests
import streamlit as st
from openpyxl import load_workbook

# ── Page config ───────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="FLARE PWR Simulator · Reactor Safety Analysis",
    page_icon="⚛",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Styling (REVISED) ─────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=IBM+Plex+Mono:wght@400;600&family=IBM+Plex+Sans:wght@400;500;600&display=swap');

:root {
    --bg:       #f5f7fa;
    --surface:  #ffffff;
    --sidebar:  #1a1f2e;
    --border:   #d0d7de;
    --accent:   #0969da;
    --accent2:  #1a7f37;
    --warn:     #9a6700;
    --danger:   #cf222e;
    --text:     #1f2328;
    --muted:    #57606a;

    --sid-text: #e6edf3;
    --sid-muted:#8b949e;

    --dark-surface: #0d1117;
    --dark-border:  #30363d;
}

html, body, [class*="css"] {
    font-family: 'IBM Plex Sans', sans-serif;
    color: var(--text);
}
.stApp { background: var(--bg); }
header { background: transparent !important; }

section[data-testid="stSidebar"] {
    background: var(--sidebar) !important;
    border-right: 1px solid var(--dark-border);
    color: var(--sid-text);
}

section[data-testid="stSidebar"] p,
section[data-testid="stSidebar"] span,
section[data-testid="stSidebar"] label {
    color: var(--sid-text) !important;
}

section[data-testid="stSidebar"] .stCaption {
    color: var(--sid-muted) !important;
}

section[data-testid="stSidebar"] .stSelectbox div[data-baseweb="select"] {
    background: var(--dark-surface);
    border: 1px solid var(--dark-border);
    border-radius: 6px;
}
section[data-testid="stSidebar"] .stSelectbox span {
    color: #ffffff !important;
}

div[role="listbox"] {
    background: var(--dark-surface) !important;
}
div[role="option"] {
    color: #ffffff !important;
}
div[role="option"]:hover {
    background: #21262d !important;
}

section[data-testid="stSidebar"] input {
    background: var(--dark-surface) !important;
    color: #ffffff !important;
    border: 1px solid var(--dark-border) !important;
}

section[data-testid="stSidebar"] button[kind="primary"] {
    background: var(--danger) !important;
    color: white !important;
}
section[data-testid="stSidebar"] button[kind="primary"]:hover {
    background: #a40e26 !important;
}

section[data-testid="stSidebar"] code {
    color: #ffffff !important;
    background: var(--dark-border) !important;
    padding: 2px 6px;
    border-radius: 4px;
}

/* Keep expander header dark when expanded so label stays visible */
section[data-testid="stSidebar"] details,
section[data-testid="stSidebar"] details[open],
section[data-testid="stSidebar"] [data-testid="stExpander"],
section[data-testid="stSidebar"] [data-testid="stExpander"] > details {
    background: var(--dark-surface) !important;
    border: 1px solid var(--dark-border) !important;
    border-radius: 4px;
}
section[data-testid="stSidebar"] details summary,
section[data-testid="stSidebar"] details[open] summary,
section[data-testid="stSidebar"] [data-testid="stExpander"] summary {
    background: var(--dark-surface) !important;
    color: var(--sid-text) !important;
}
section[data-testid="stSidebar"] details summary p,
section[data-testid="stSidebar"] details summary span,
section[data-testid="stSidebar"] details[open] summary p,
section[data-testid="stSidebar"] details[open] summary span {
    color: var(--sid-text) !important;
}

h1, h2, h3 {
    font-family: 'IBM Plex Mono', monospace;
}

button[role="tab"] {
    color: var(--muted);
}
button[role="tab"][aria-selected="true"] {
    color: var(--accent);
    border-bottom: 2px solid var(--accent);
}

.console {
    background: #0d1117;
    border: 1px solid #30363d;
    border-radius: 6px;
    padding: 1rem;
    color: #7ee787;
    font-family: 'IBM Plex Mono', monospace;
}

.metric-grid {
    display:grid;
    grid-template-columns:repeat(4,1fr);
    gap:0.75rem;
}
.metric-tile {
    background: var(--surface);
    border: 1px solid var(--border);
    border-radius: 6px;
    padding: 1rem;
    text-align: center;
}
.metric-tile .val {
    font-family: 'IBM Plex Mono', monospace;
    font-size: 1.5rem;
    color: var(--accent);
}
.metric-tile .val.small {
    font-size: 0.82rem;
    line-height: 1.4;
}
.metric-tile .lbl.small {
    font-size: 0.68rem;
}
.metric-tile.danger .val { color: var(--danger); }
.metric-tile.warn   .val { color: var(--warn); }
.metric-tile.ok     .val { color: var(--accent2); }

.hdiv {
    border-top: 1px solid var(--border);
    margin: 1.5rem 0;
}
/* ── Tooltip (?) icon — targets the actual SVG stroke ── */
section[data-testid="stSidebar"] .stTooltipHoverTarget svg.icon {
    stroke: #f97316 !important;
}
section[data-testid="stSidebar"] .stTooltipHoverTarget:hover svg.icon {
    stroke: #ffffff !important;
}
.stTooltipHoverTarget svg.icon {
    stroke: #0969da !important;
}
</style>
""", unsafe_allow_html=True)

# ── Constants ──────────────────────────────────────────────────────────────────

# When exec'd by flare_home.py, FLARE_WORK_DIR is injected into globals.
# Fall back to __file__ when running standalone.
WORK_DIR = Path(globals().get("FLARE_WORK_DIR", Path(__file__).parent))

def _runtime_dir(base_dir: Path | None = None) -> Path:
    """Return the FLARE runtime folder, preferring runtime/ then Runtime/."""
    base = Path(base_dir) if base_dir is not None else WORK_DIR
    lower = base / "runtime"
    upper = base / "Runtime"
    if lower.exists():
        rt = lower
    elif upper.exists():
        rt = upper
    else:
        rt = lower
        rt.mkdir(parents=True, exist_ok=True)
    return rt


def _runtime_file(name: str, base_dir: Path | None = None) -> Path:
    return _runtime_dir(base_dir) / name


# ── Anthropic API key (shared with flare_home.py) ─────────────────────────────
def _read_config(key):
    """Read a key=value entry from flare_config.txt.

    Accepts whitespace around the equals sign, ignores comments/blank lines,
    and compares keys case-insensitively so a harmless formatting change in
    flare_config.txt does not make the UI think the API key is missing.
    """
    cfg = _runtime_file("flare_config.txt")
    if cfg.exists():
        try:
            for line in cfg.read_text(encoding="utf-8").splitlines():
                line = line.strip()
                if not line or line.startswith("#") or "=" not in line:
                    continue
                k, v = line.split("=", 1)
                if k.strip().lower() == key.lower():
                    return v.strip().strip('"').strip("'")
        except Exception:
            return None
    return None

def load_api_key():
    # Priority order:
    #   1. flare_config.txt  (persistent local FLARE setting)
    #   2. st.session_state.user_api_key  (entered on flare_home during this session)
    #   3. ANTHROPIC_API_KEY environment variable
    # This makes the Simulator narrative recognize a key entered on the home
    # page even before the user saves it, while still supporting persistent
    # config-file storage.
    _cfg = _read_config("ANTHROPIC_API_KEY")
    if _cfg:
        return _cfg
    try:
        _sess = st.session_state.get("user_api_key", "")
        if _sess:
            return _sess
    except Exception:
        pass
    return os.environ.get("ANTHROPIC_API_KEY")

def load_model():
    """Read ANTHROPIC_MODEL from flare_config.txt; fall back to claude-sonnet-4-5."""
    return _read_config("ANTHROPIC_MODEL") or "claude-sonnet-4-5"

_API_KEY         = load_api_key()
_ANTHROPIC_MODEL = load_model()


NARRATIVE_QUALITY_GUIDANCE = """
AI narrative quality requirements:
- Use a tiered scope. For verification, null, or benign cases with no meaningful excursion, use only three concise sections: Steady-State/Initial Condition Performance, Code Behavior, and Regulatory Significance. Do not force full accident-analysis boilerplate onto benign cases.
- Reserve detailed Fuel Integrity and Core Damage and Source Term and Radiological Consequences sections for cases with nonzero ECR, hydrogen generation, rod failures, radionuclide releases, or nonzero dose.
- If all radionuclide group releases are zero, state that in one sentence and do not enumerate radionuclide groups individually.
- If ECR and hydrogen generation are zero, state that in one sentence and do not expand into a full oxidation discussion.
- The 10 CFR 50.46 PCT limit is 2,200°F (1,204°C / 1,477 K). State clearly whether the result passes or fails this criterion. If peak hot-pin clad temperature exceeds 1,477 K, call it a limit exceedance; do not hedge with phrases such as 'slightly exceeded' while also implying compliance.
- Describe deterministic code behavior declaratively. Avoid hedging phrases such as 'likely', 'suggests', 'may indicate', or 'appears to be' when explaining model responses already contained in the output.
- DNBR equal to NaN indicates that the DNB correlation is not applicable under the final conditions, not a numerical failure. Acknowledge it in one sentence only if needed.
- If initial SG heat removal is less than 50% of initial reactor power, state that the case was initialized at partial-load, hot-zero-power, or otherwise non-full-power conditions; do not call it a normal full-power case.
- If any dose receptor exceeds its regulatory limit, state this explicitly in conclusions, not only in tabulated data.
- Do not reproduce source-term or dose tables in prose. The numerical tables are appended separately; discuss only the important implications.
"""

PLOT_COLOR = "#0969da"
PLOT_BG    = "#ffffff"
PLOT_PAPER = "#f5f7fa"
PLOT_GRID  = "#d0d7de"
PLOT_TEXT  = "#1f2328"

# Unit conversion table  (matches flare_ua.py)
# suffix → (si_label, eng_label, si_fn, eng_fn, eng_scale)
UNIT_CONV = {
    "(K)":    ("°C",   "°F",   lambda v: v - 273.15,        lambda v: (v - 273.15)*9/5 + 32, 9/5),
    "(kPa)":  ("kPa",  "psia", lambda v: v,                  lambda v: v * 0.145038,           0.145038),
    "(kg/s)": ("kg/s", "lb/s", lambda v: v,                  lambda v: v * 2.20462,            2.20462),
    "(MW)":   ("MW",   "MW",   lambda v: v,                  lambda v: v,                      1.0),
    "(m/s)":  ("m/s",  "ft/s", lambda v: v,                  lambda v: v * 3.28084,            3.28084),
    "(rpm)":  ("rpm",  "rpm",  lambda v: v,                  lambda v: v,                      1.0),
}


# ── Helpers ───────────────────────────────────────────────────────────────────

def smart_step(v):
    """
    Return a step size equal to the place value of the lowest significant digit.

    Examples:
        100    → 100      12340  → 10
        101    → 1        1.455  → 0.001
        0.05   → 0.01     4000   → 1000
        5200   → 100
    """
    if v == 0:
        return 0.01
    av = abs(float(v))
    s = f"{av:.10g}"                        # up to 10 sig-figs, no trailing zeros

    if 'e' in s or 'E' in s:               # very large/small → use magnitude
        return 10 ** math.floor(math.log10(av))

    if '.' in s:
        int_part, dec_part = s.split('.')
        # last non-zero decimal position (0-indexed from decimal point)
        for i in range(len(dec_part) - 1, -1, -1):
            if dec_part[i] != '0':
                return 10 ** (-(i + 1))
        s = int_part                        # all decimal digits were zero

    # integer part: find last non-zero digit scanning right-to-left
    for i, ch in enumerate(reversed(s)):
        if ch != '0':
            return float(10 ** i)
    return 1.0


# ── Input discovery ───────────────────────────────────────────────────────────
# User-maintained input decks now live in subfolders below the FLARE root.
# Generated output/control folders are excluded, and root-level *_in.xlsx files
# are intentionally ignored.
_EXCLUDED_INPUT_DIR_PREFIXES = ("sim_", "risk_", "ua_", "sim_all_", ".sim_all_", "__pycache__")

def _is_generated_input_dir(path: Path) -> bool:
    try:
        rel_parts = path.relative_to(WORK_DIR).parts
    except Exception:
        rel_parts = path.parts
    for part in rel_parts:
        if part.startswith(_EXCLUDED_INPUT_DIR_PREFIXES) or part.startswith("."):
            return True
    return False

def discover_input_cases():
    entries = []
    for p in WORK_DIR.rglob("*_in.xlsx"):
        if p.parent == WORK_DIR:
            continue  # root-level input files are intentionally ignored
        if p.stem.startswith(".~") or p.stem.startswith("ua_"):
            continue
        if _is_generated_input_dir(p.parent):
            continue
        case = p.stem[:-3] if p.stem.endswith("_in") else p.stem.replace("_in", "")
        rel = p.parent.relative_to(WORK_DIR).as_posix()
        label = f"{case}  —  {rel}"
        entries.append({"case": case, "path": p, "rel": rel, "label": label})
    entries.sort(key=lambda e: (e["case"].lower(), e["rel"].lower()))
    return entries

def find_cases():
    return [e["label"] for e in discover_input_cases()]

def _case_from_label(label: str):
    for e in discover_input_cases():
        if e["label"] == label:
            return e
    return None


def load_command_block(case_name, input_path=None):
    """Return list of (row, text) for all command rows.

    The UI may call this while the user's workbook is open in Excel.  Do not
    open the original workbook directly.  Take a private snapshot first, then
    read the snapshot.  If Excel/OneDrive has the file momentarily locked, fail
    soft so the simulator page still loads.
    """
    p = Path(input_path) if input_path is not None else WORK_DIR / f"{case_name}_in.xlsx"
    if not p.exists():
        return []

    import tempfile as _tempfile
    import uuid as _uuid

    tmp = Path(_tempfile.gettempdir()) / f"FLARE_UI_{case_name}_{os.getpid()}_{_uuid.uuid4().hex[:8]}_in.xlsx"
    wb = None
    try:
        shutil.copy2(p, tmp)
        wb = load_workbook(tmp, read_only=True, data_only=True)
        ws = wb[f"{case_name}_in"]
        rows = []
        for i, row in enumerate(ws.iter_rows(max_col=1, values_only=True), 1):
            v = row[0]
            if isinstance(v, (int, float)) and not isinstance(v, bool):
                break
            if isinstance(v, str):
                rows.append((i, v))
            if i > 120:
                break
        return rows
    except PermissionError:
        # Excel/OneDrive occasionally locks the file during save/sync.  The
        # sidebar command preview is nonessential; the simulation engine will
        # take its own snapshot at run time.
        return []
    except Exception:
        return []
    finally:
        try:
            if wb is not None:
                wb.close()
        except Exception:
            pass
        try:
            tmp.unlink(missing_ok=True)
        except Exception:
            pass


def parse_params(rows):
    """Extract {name: value} from command block rows."""
    params = {}
    for _, text in rows:
        if text.strip().startswith("#"):
            continue
        m = re.match(r"^\s*(\w+)\s*=\s*([+-]?\d+\.?\d*(?:[eE][+-]?\d+)?)", text)
        if m:
            try:
                params[m.group(1)] = float(m.group(2))
            except ValueError:
                pass
    return params


def build_override_xlsx(case_name, overrides: dict, input_path=None):
    """
    Copy base _in.xlsx to a temp file (.~run_in.xlsx) applying overrides,
    then set minimum_output=0 so full post-processing runs.
    Returns temp case name.

    The original workbook may be open in Excel, so copy it first and only open
    the private copy with openpyxl.
    """
    src = Path(input_path) if input_path is not None else WORK_DIR / f"{case_name}_in.xlsx"
    tmp_name = f".~{case_name}_run"
    dst = WORK_DIR / f"{tmp_name}_in.xlsx"

    shutil.copy2(src, dst)

    wb = load_workbook(dst)
    ws = wb[f"{case_name}_in"]
    ws.title = f"{tmp_name}_in"
    if f"{case_name}_out" in wb.sheetnames:
        wb[f"{case_name}_out"].title = f"{tmp_name}_out"

    for var, new_val in overrides.items():
        pat = re.compile(
            r"^(\s*" + re.escape(var) + r"\s*=\s*)"
            r"([+-]?\d+\.?\d*(?:[eE][+-]?\d+)?)"
            r"(.*)"
        )
        for row in ws.iter_rows(max_col=1):
            cell = row[0]
            if isinstance(cell.value, str) and var in cell.value and "=" in cell.value:
                m = pat.match(cell.value)
                if m:
                    vs = (f"{new_val:.6g}" if abs(new_val) < 1e4
                          else f"{new_val:.4e}")
                    cell.value = m.group(1) + vs + m.group(3)
                    break

    wb.save(dst)
    wb.close()
    return tmp_name


def run_simulation(case_name, progress_queue, no_figures=False):
    """Run simulation in a thread, pushing stdout lines to queue."""
    cmd = [sys.executable, str(WORK_DIR / "flare_sim.py"), case_name]
    if no_figures:
        cmd.append("--no-figures")
    proc = subprocess.Popen(
        cmd,
        stdout=subprocess.PIPE,
        stderr=subprocess.STDOUT,
        text=True,
        cwd=str(WORK_DIR),
    )
    for line in proc.stdout:
        progress_queue.put(("line", line.rstrip()))
    proc.wait()
    progress_queue.put(("done", proc.returncode))


def load_results(case_name, run_dir=None):
    """Load _out.xlsx Sheet1 into a DataFrame.  Looks in run_dir first."""
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        p = Path(search_dir) / f"{case_name}_out.xlsx"
        if p.exists():
            try:
                return pd.read_excel(p, sheet_name="Sheet1", engine="openpyxl")
            except Exception:
                return None
    return None


def load_diagnostics(case_name, run_dir=None):
    """Load the separate FLARE diagnostics CSV if it exists.

    The SG/pump diagnostic stream is written by flare_sim.py as
    <case>_diag.csv in the active run directory.  Results display should be
    tolerant of older output folders that do not have this file.
    """
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        if search_dir is None:
            continue
        p = Path(search_dir) / f"{case_name}_diag.csv"
        if p.exists():
            try:
                return pd.read_csv(p)
            except Exception:
                return None
    return None


# ── Durable "Run All Cases" batch-worker support ─────────────────────────────
def _sim_all_read_json(path, default=None):
    try:
        p = Path(path)
        if p.exists():
            return json.loads(p.read_text(encoding="utf-8"))
    except Exception:
        pass
    return {} if default is None else default


def _sim_all_write_json_atomic(path, payload):
    """Best-effort JSON write for UI-side control files."""
    try:
        p = Path(path)
        p.parent.mkdir(parents=True, exist_ok=True)
        tmp = p.with_name(f"{p.name}.{os.getpid()}.tmp")
        tmp.write_text(json.dumps(payload, indent=2), encoding="utf-8")
        try:
            os.replace(str(tmp), str(p))
        except Exception:
            p.write_text(json.dumps(payload, indent=2), encoding="utf-8")
            try:
                tmp.unlink(missing_ok=True)
            except Exception:
                pass
    except Exception:
        pass


def _sim_all_tail_text_file(path, max_chars=8000):
    try:
        p = Path(path)
        if not p.exists():
            return ""
        with open(p, "rb") as f:
            try:
                f.seek(0, 2)
                size = f.tell()
                f.seek(max(size - max_chars, 0))
            except Exception:
                pass
            data = f.read()
        return data.decode("utf-8", errors="replace")
    except Exception as e:
        return f"(Could not read log: {e})"


def _sim_all_latest_status_dir():
    try:
        dirs = sorted(
            [d for d in WORK_DIR.iterdir()
             if d.is_dir() and (d.name.startswith(".sim_all_") or d.name.startswith("sim_all_"))],
            key=lambda d: d.stat().st_mtime,
            reverse=True,
        )
        for d in dirs:
            if (d / "sim_all_status.json").exists():
                return d
    except Exception:
        pass
    return None


def _sim_all_terminate_process_tree(pid):
    try:
        pid = int(pid)
    except Exception:
        return False, "No valid worker PID was available."
    try:
        if sys.platform.startswith("win"):
            cp = subprocess.run(
                ["taskkill", "/PID", str(pid), "/T", "/F"],
                capture_output=True, text=True, check=False,
            )
            msg = (cp.stdout or cp.stderr or "").strip()
            return cp.returncode == 0, msg or f"taskkill returned {cp.returncode}"
        else:
            import signal as _signal
            try:
                os.killpg(pid, _signal.SIGTERM)
            except Exception:
                os.kill(pid, _signal.SIGTERM)
            return True, "Terminate signal sent."
    except Exception as e:
        return False, str(e)


def _sim_all_abort(run_dir):
    run_dir = Path(run_dir)
    status_path = run_dir / "sim_all_status.json"
    status = _sim_all_read_json(status_path, {})
    _sim_all_write_json_atomic(run_dir / "sim_all_abort_requested.json", {
        "abort_requested": True,
        "requested_at": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        "reason": "User clicked Abort Run All Cases in the FLARE UI.",
    })
    pid = status.get("worker_pid") or status.get("pid")
    ok, msg = (False, "No worker PID found in sim_all_status.json.")
    if pid:
        ok, msg = _sim_all_terminate_process_tree(pid)
    status.update({
        "status": "aborted",
        "message": "Batch run aborted by user." if ok else f"Abort requested; manual process cleanup may be needed: {msg}",
        "abort_requested": True,
        "last_update": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
    })
    _sim_all_write_json_atomic(status_path, status)
    return ok, msg


def _sim_all_launch_worker(run_dir, case_list, fast_mode, final_report=False, final_report_detail=0.5):
    worker = WORK_DIR / "flare_sim_batch_worker.py"
    if not worker.exists():
        st.error(
            "`flare_sim_batch_worker.py` was not found in the FLARE folder. "
            "Place the worker file alongside `flare_ui.py`."
        )
        return False

    run_dir = Path(run_dir)
    run_dir.mkdir(parents=True, exist_ok=True)
    # run_dir is only a small regular control/status folder.
    # The worker writes each case output to a normal sim_<case>_<timestamp> folder.
    batch_tag = run_dir.name.replace(".sim_all_", "").replace("sim_all_", "")
    # Persist both case names and input paths. Older workers accept a simple
    # list of names; the recursive-input worker uses case_entries.
    _case_entries_cfg = []
    for item in case_list:
        if isinstance(item, dict):
            _case_entries_cfg.append({
                "case": item["case"],
                "input_path": str(item["path"]),
                "rel": item.get("rel", ""),
            })
        else:
            _case_entries_cfg.append({"case": str(item), "input_path": str(WORK_DIR / f"{item}_in.xlsx"), "rel": ""})
    cfg = {
        "work_dir": str(WORK_DIR),
        "run_dir": str(run_dir),
        "batch_tag": batch_tag,
        "output_mode": "per_case_sim_dirs",
        "cases": [e["case"] for e in _case_entries_cfg],
        "case_entries": _case_entries_cfg,
        "fast_mode": bool(False if final_report else fast_mode),
        "final_report": bool(final_report),
        "final_report_detail": float(final_report_detail),
        "timeout_s": 600,
    }
    cfg_path = run_dir / "sim_all_worker_config.json"
    _sim_all_write_json_atomic(cfg_path, cfg)
    try:
        (run_dir / "sim_all_abort_requested.json").unlink()
    except Exception:
        pass

    status = {
        "status": "starting",
        "message": "Starting batch simulation worker…",
        "total_runs": len(_case_entries_cfg),
        "completed_runs": 0,
        "failed_runs": 0,
        "current_case": None,
        "run_dir": str(run_dir),
        "last_update": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
    }
    _sim_all_write_json_atomic(run_dir / "sim_all_status.json", status)

    stdout_f = open(run_dir / "sim_all_stdout.log", "ab", buffering=0)
    stderr_f = open(run_dir / "sim_all_stderr.log", "ab", buffering=0)
    creationflags = getattr(subprocess, "CREATE_NEW_PROCESS_GROUP", 0) if sys.platform.startswith("win") else 0
    env = {**os.environ, "PYTHONUTF8": "1", "MPLBACKEND": "Agg"}
    try:
        proc = subprocess.Popen(
            [sys.executable, str(worker), str(cfg_path)],
            cwd=str(WORK_DIR),
            stdout=stdout_f,
            stderr=stderr_f,
            stdin=subprocess.DEVNULL,
            env=env,
            creationflags=creationflags,
            close_fds=(not sys.platform.startswith("win")),
        )
        status.update({
            "status": "running",
            "message": "Batch simulation worker launched.",
            "worker_pid": int(proc.pid),
            "last_update": datetime.now().strftime("%Y-%m-%dT%H:%M:%S"),
        })
        _sim_all_write_json_atomic(run_dir / "sim_all_status.json", status)
        st.session_state.sim_all_run_dir = run_dir
        return True
    except Exception as e:
        st.error(f"Could not launch batch simulation worker: {e}")
        return False


def _sim_all_load_status(run_dir=None):
    if run_dir is None:
        run_dir = st.session_state.get("sim_all_run_dir")
    if not run_dir:
        run_dir = _sim_all_latest_status_dir()
        if run_dir:
            st.session_state.sim_all_run_dir = run_dir
    if not run_dir:
        return {}
    return _sim_all_read_json(Path(run_dir) / "sim_all_status.json", {})


def _sim_all_render_progress(status):
    if not status:
        return
    state = str(status.get("status", "unknown"))
    total = max(int(status.get("total_runs", 0) or 0), 1)
    done = int(status.get("completed_runs", 0) or 0)
    failed = int(status.get("failed_runs", 0) or 0)
    frac = min(max(done / total, 0.0), 1.0)
    run_dir = Path(status.get("run_dir", "")) if status.get("run_dir") else None
    msg = status.get("message", "")

    st.markdown("### Run All Cases")
    if state in {"starting", "running"}:
        st.progress(frac, text=f"{done}/{total} complete — {msg}")
        c1, c2, c3, c4 = st.columns(4)
        c1.metric("Completed", f"{done} / {total}")
        c2.metric("Failed", str(failed))
        c3.metric("Current case", status.get("current_case") or "—")
        elapsed = status.get("current_case_elapsed_s")
        c4.metric("Case elapsed", f"{elapsed:.0f} s" if isinstance(elapsed, (int, float)) else "—")
        if run_dir:
            st.caption(f"Batch status folder: `{run_dir.name}`; outputs are being written to ordinary `sim_<case>_<timestamp>` folders.")
            if status.get("final_report"):
                st.caption("Final Report mode is active: figures are forced on and a final PDF will be compiled at the end.")
        log_path = status.get("current_case_console_log")
        if log_path:
            with st.expander("Current case console log", expanded=True):
                st.code(_sim_all_tail_text_file(log_path, max_chars=10000) or "(No console output yet.)", language="text")
        if st.button("⛔ Abort Run All Cases", type="primary", key="sim_all_abort_btn"):
            ok, abort_msg = _sim_all_abort(run_dir)
            if ok:
                st.warning("Batch run aborted. The worker process tree was terminated.")
            else:
                st.error(f"Abort requested, but process termination could not be confirmed: {abort_msg}")
            st.rerun()
        time.sleep(1.0)
        st.rerun()
    elif state == "complete":
        st.success(f"Batch run complete — {done - failed}/{total} OK, {failed} failed.")
        if run_dir:
            st.caption(f"Batch status folder: `{run_dir.name}`; outputs are being written to ordinary `sim_<case>_<timestamp>` folders.")
            summary_csv = run_dir / "FLARE_sim_all_results.csv"
            if summary_csv.exists():
                with open(summary_csv, "rb") as f:
                    st.download_button("📥 Download batch results CSV", data=f.read(), file_name=summary_csv.name, mime="text/csv")
            final_pdf = Path(status.get("final_report_pdf") or (run_dir / "FLARE_Run_All_Final_Report.pdf"))
            if final_pdf.exists():
                with open(final_pdf, "rb") as f:
                    st.download_button(
                        "📥 Download Run All Final Report PDF",
                        data=f.read(),
                        file_name=final_pdf.name,
                        mime="application/pdf",
                    )
    elif state == "aborted":
        st.warning(f"Batch run aborted — {done}/{total} cases processed.")
        if run_dir:
            st.caption(f"Batch status folder: `{run_dir.name}`; outputs are being written to ordinary `sim_<case>_<timestamp>` folders.")
    elif state in {"failed", "error"}:
        st.error(f"Batch run failed — {msg}")
        if run_dir:
            err_log = run_dir / "sim_all_stderr.log"
            if err_log.exists():
                with st.expander("Worker error log", expanded=True):
                    st.code(_sim_all_tail_text_file(err_log), language="text")


def load_source_term(case_name, run_dir=None):
    """Load Source Term sheet if it exists.  Looks in run_dir first."""
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        p = Path(search_dir) / f"{case_name}_out.xlsx"
        if p.exists():
            try:
                wb = load_workbook(p, read_only=True)
                if "Source Term" not in wb.sheetnames:
                    wb.close()
                    return None
                df = pd.read_excel(p, sheet_name="Source Term", engine="openpyxl")
                wb.close()
                return df
            except Exception:
                return None
    return None



def load_input_echo_value(case_name, variable_name, run_dir=None):
    """Return a value from the output workbook's Input Echo sheet, if present."""
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        p = Path(search_dir) / f"{case_name}_out.xlsx"
        if not p.exists():
            continue
        try:
            wb = load_workbook(p, read_only=True, data_only=True)
            if "Input Echo" not in wb.sheetnames:
                wb.close()
                continue
            ws = wb["Input Echo"]
            rows = list(ws.iter_rows(values_only=True))
            wb.close()
            if not rows:
                continue
            header = [str(x).strip() if x is not None else "" for x in rows[0]]
            try:
                ivar = header.index("Variable")
                ival = header.index("Value")
            except ValueError:
                ivar, ival = 0, 1
            for row in rows[1:]:
                if len(row) <= max(ivar, ival):
                    continue
                if str(row[ivar]).strip().lower() == variable_name.lower():
                    return "" if row[ival] is None else str(row[ival]).strip()
        except Exception:
            continue
    return None


def _format_source_term_option(value):
    """Return a user-facing spelling for source-term options."""
    if value is None:
        return None
    raw = str(value).strip().strip('"').strip("'")
    if not raw or raw.lower() in {"nan", "none"}:
        return None
    norm = raw.lower().replace("_", "-")
    aliases = {
        "licensing-auto": "licensing-auto",
        "licensing auto": "licensing-auto",
        "thermal-failure": "thermal-failure",
        "rg1183-loca": "RG1183-LOCA",
        "rg-1183-loca": "RG1183-LOCA",
        "rg1183-nonloca": "RG1183-nonLOCA",
        "rg1183-non-loca": "RG1183-nonLOCA",
        "rg-1183-nonloca": "RG1183-nonLOCA",
        "rg1183-nonloca-dnb": "RG1183-nonLOCA-DNB",
        "rg1183-non-loca-dnb": "RG1183-nonLOCA-DNB",
    }
    return aliases.get(norm, raw.replace("_", "-"))


def _extract_source_term_basis(st_df, case_name=None, run_dir=None):
    """Extract source-term option/applied model from the Source Term sheet.

    Newer FLARE outputs include metadata rows in the Source Term sheet:
    'Fuel source-term option' is the user-selected option and 'Model selected'
    is the licensing-auto resolution / applied reduced-order model.  For older
    files, fall back to Input Echo if available.
    """
    option = None
    applied = None
    if st_df is not None and not st_df.empty and "Group" in st_df.columns:
        value_cols = [c for c in ["Total release frac", "Value", "Model"] if c in st_df.columns]
        if not value_cols:
            value_cols = [c for c in st_df.columns if c != "Group"][:1]
        for _, row in st_df.iterrows():
            label = str(row.get("Group", "")).strip().lower()
            val = None
            for c in value_cols:
                vv = row.get(c)
                if vv is not None and not pd.isna(vv) and str(vv).strip() != "":
                    val = vv
                    break
            if label in {"fuel source-term option", "source term option", "source_term_model"}:
                option = val
            elif label in {"model selected", "applied source-term model", "applied model"}:
                applied = val
    if option is None and case_name is not None:
        option = load_input_echo_value(case_name, "source_term_model", run_dir=run_dir)
    return _format_source_term_option(option), _format_source_term_option(applied)

def load_iodine_spike(case_name, run_dir=None):
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        p = Path(search_dir) / f"{case_name}_out.xlsx"
        if p.exists():
            try:
                wb = load_workbook(p, read_only=True)
                if "Iodine Spike" not in wb.sheetnames:
                    wb.close()
                    return None
                ws = wb["Iodine Spike"]
                rows = [[c.value for c in row] for row in ws.iter_rows()]
                wb.close()
                return rows
            except Exception:
                return None
    return None




def _clean_numeric_columns_for_streamlit(df, numeric_cols=None, text_cols=None):
    """Return a copy with stable dtypes before st.dataframe/Arrow serialization.

    Streamlit serializes displayed dataframes through PyArrow.  Columns read
    from Excel can arrive as object dtype with a mix of ints, floats, blanks,
    and strings; PyArrow may then infer a string/bytes column and fail when it
    encounters a number.  Convert known numeric columns explicitly and known
    label/result columns to strings.
    """
    out = df.copy()
    numeric_cols = numeric_cols or []
    text_cols = text_cols or []
    for col in numeric_cols:
        if col in out.columns:
            out[col] = pd.to_numeric(out[col], errors="coerce")
    for col in text_cols:
        if col in out.columns:
            out[col] = out[col].fillna("").astype(str)
    return out


def _clean_dose_display_df(df):
    """Normalize common NOTBADTRAD/Iodine Spike display columns."""
    numeric_cols = [
        "TEDE (rem)", "Margin (rem)",
        "Margin (× limit)", "Released (Ci)", "EAB (rem)",
        "LPZ (rem)", "CR (rem)", "Release frac",
        "Distance (m)", "χ/Q (s/m³)", "chi/Q (s/m3)",
        "TEDE EAB interval (rem)", "TEDE LPZ interval (rem)",
        "Total release frac",
    ]
    # Keep Limit as text. The workbook intentionally stores entries such as
    # "25 rem (2 hr)" / "5 rem". Converting this column to numeric makes
    # Streamlit display it as None/NaN even though the regulatory limits are known.
    text_cols = ["Location", "Result", "Group", "Limit (rem)"]
    return _clean_numeric_columns_for_streamlit(df, numeric_cols, text_cols)



def _clean_source_term_display_df(df, numeric_cols=None, text_cols=None):
    """Normalize Source Term display columns before Streamlit/PyArrow rendering."""
    default_numeric = [
        "Core inventory (Ci)",
        "Released inventory (Ci)",
        "Gap release %",
        "Early IV %",
        "Total release %",
        "Total release frac",
        "NOTBADTRAD gap release duration (hr)",
        "NOTBADTRAD early-IV release duration (hr)",
        "NOTBADTRAD EAB integration time (hr)",
        "NOTBADTRAD LPZ integration time (hr)",
        "Estimated fissions",
    ]
    default_text = ["Group", "NBT key", "Model selected", "Fuel source-term option", "Inventory model"]
    return _clean_numeric_columns_for_streamlit(
        df,
        (numeric_cols or default_numeric),
        (text_cols or default_text),
    )



def _source_term_lines_for_narrative(case_name, run_dir=None, max_groups=12):
    """Return compact, narrative-ready source-term and dose lines.

    The AI narrative uses these lines to keep fuel-integrity discussion
    separate from source-term/radiological consequences while still grounding
    the source-term section in the same data shown in the Results panel.
    """
    lines = []
    try:
        st_df = load_source_term(case_name, run_dir=run_dir)
        if st_df is not None and not st_df.empty:
            option, applied = _extract_source_term_basis(st_df, case_name=case_name, run_dir=run_dir)
            if option:
                if option == "licensing-auto":
                    lines.append(f"Source-term option: licensing-auto; applied model: {applied or 'not reported'}")
                else:
                    lines.append(f"Source-term option: {option}")

            inv_col = "Core inventory (Ci)"
            rel_col = "Released inventory (Ci)"
            frac_col = "Total release frac"
            if inv_col in st_df.columns and "Group" in st_df.columns:
                df = st_df.copy()
                for c in [inv_col, rel_col, frac_col, "Gap release %", "Early IV %", "Total release %", "BDBE F factor", "BDBE ECR adjustment %"]:
                    if c in df.columns:
                        df[c] = pd.to_numeric(df[c], errors="coerce")
                main = df[df[inv_col].notna()].copy()
                if not main.empty:
                    total_rows = main[main["Group"].astype(str).str.upper().eq("TOTAL")]
                    non_total = main[~main["Group"].astype(str).str.upper().eq("TOTAL")].copy()
                    # Suppress boilerplate for benign cases: if every group release is zero,
                    # tell the AI that once and do not feed it nine zero-release rows.
                    _release_cols = [c for c in [rel_col, frac_col, "Total release %"] if c in non_total.columns]
                    _all_zero_release = False
                    if _release_cols:
                        _vals = []
                        for _c in _release_cols:
                            _vals.extend(pd.to_numeric(non_total[_c], errors="coerce").fillna(0.0).abs().tolist())
                        _all_zero_release = bool(_vals) and max(_vals) <= 0.0
                    if _all_zero_release:
                        lines.append("All radionuclide group releases are zero; source-term prose should state this once and should not enumerate individual groups.")
                    else:
                        if rel_col in non_total.columns:
                            non_total = non_total.sort_values(rel_col, ascending=False)
                        rows = list(non_total.head(max_groups).to_dict("records"))
                        if not total_rows.empty:
                            rows.append(total_rows.iloc[0].to_dict())
                        lines.append("Source-term group releases shown in the Results panel:")
                        for r in rows:
                            g = str(r.get("Group", "")).strip()
                            if not g or g.lower() == "nan":
                                continue
                            parts = [g]
                            if "Total release %" in r and pd.notna(r.get("Total release %")):
                                parts.append(f"total release {float(r['Total release %']):.4g}%")
                            elif frac_col in r and pd.notna(r.get(frac_col)):
                                parts.append(f"total release fraction {float(r[frac_col]):.4e}")
                            if rel_col in r and pd.notna(r.get(rel_col)):
                                parts.append(f"released inventory {float(r[rel_col]):.4e} Ci")
                            if "Gap release %" in r and pd.notna(r.get("Gap release %")):
                                parts.append(f"gap {float(r['Gap release %']):.4g}%")
                            if "Early IV %" in r and pd.notna(r.get("Early IV %")):
                                parts.append(f"early-IV {float(r['Early IV %']):.4g}%")
                            if "BDBE F factor" in r and pd.notna(r.get("BDBE F factor")):
                                parts.append(f"BDBE F {float(r['BDBE F factor']):.4g}")
                            if "BDBE ECR adjustment %" in r and pd.notna(r.get("BDBE ECR adjustment %")):
                                parts.append(f"ECR adjustment {float(r['BDBE ECR adjustment %']):.4g}%")
                            lines.append(" - " + "; ".join(parts))

            # Metadata rows displayed in Results panel as source-term basis.
            if "Group" in st_df.columns:
                meta_labels = {
                    "Model selected", "Fuel source-term option", "Inventory model",
                    "Estimated fissions", "Severe-event ECR release flag",
                    "Severe-event ECR threshold (%)", "Predicted mean oxidizing-rod ECR (%)",
                    "Severe-event ECR active",
                    "NOTBADTRAD gap release duration (hr)",
                    "NOTBADTRAD early-IV release duration (hr)",
                    "NOTBADTRAD EAB integration time (hr)",
                    "NOTBADTRAD LPZ integration time (hr)",
                }
                value_cols = [c for c in ["Total release frac", "Value", "Model", "Released inventory (Ci)"] if c in st_df.columns]
                for _, row in st_df[st_df["Group"].isin(meta_labels)].iterrows():
                    label = str(row.get("Group", "")).strip()
                    val = None
                    for c in value_cols:
                        vv = row.get(c)
                        if vv is not None and not pd.isna(vv) and str(vv).strip() != "":
                            val = vv
                            break
                    if val is not None:
                        lines.append(f"Source-term basis: {label} = {val}")

        dose_rows = load_dose(case_name, run_dir=run_dir)
        if dose_rows:
            dose_bits = []
            for r in dose_rows:
                if r and r[0] in ("EAB", "LPZ", "Control Room") and len(r) >= 5:
                    dose_bits.append(f"{r[0]} TEDE {r[1]} rem vs limit {r[2]}: {r[4]}")
            if dose_bits:
                lines.append("NOTBADTRAD dose screening: " + "; ".join(dose_bits))

        isp_rows = load_iodine_spike(case_name, run_dir=run_dir)
        if isp_rows:
            isp_bits = []
            for r in isp_rows:
                if r and r[0] in ("EAB", "LPZ", "Control Room") and len(r) >= 5:
                    isp_bits.append(f"{r[0]} iodine-spike TEDE {r[1]} rem vs limit {r[2]}: {r[4]}")
            if isp_bits:
                lines.append("Iodine spike pre-existing coolant activity dose: " + "; ".join(isp_bits))
    except Exception as e:
        lines.append(f"Source-term narrative data unavailable: {e}")
    return lines


def _fmt_report_cell(v, fmt=None):
    """Format a cell for narrative report tables without changing source data."""
    try:
        if v is None or pd.isna(v):
            return ""
    except Exception:
        if v is None:
            return ""
    if fmt:
        try:
            return fmt.format(float(v))
        except Exception:
            return str(v)
    return str(v)


def _report_df(df, fmt_map=None, text_cols=None):
    """Return a string-valued dataframe for Word/PDF report export."""
    fmt_map = fmt_map or {}
    text_cols = set(text_cols or [])
    out = df.copy()
    for c in out.columns:
        fmt = fmt_map.get(c)
        if c in text_cols:
            out[c] = out[c].map(lambda v: _fmt_report_cell(v, None))
        else:
            out[c] = out[c].map(lambda v, _fmt=fmt: _fmt_report_cell(v, _fmt))
    return out


def build_source_term_dose_report_tables(case_name, run_dir=None):
    """Build the same source-term/dose tables shown in the Results panel.

    These tables are used by the AI narrative Word report so that the report
    contains the exact source-term and dose data that the user sees in the UI,
    instead of relying only on prose summaries.
    Returns a list of dictionaries with section/title/df keys.
    """
    tables = []

    # Source Term block
    st_df = load_source_term(case_name, run_dir=run_dir)
    if st_df is not None and not st_df.empty:
        try:
            option, applied = _extract_source_term_basis(st_df, case_name=case_name, run_dir=run_dir)
            basis_rows = []
            if option:
                if option == "licensing-auto":
                    basis_rows.append(["Source-term option", option])
                    basis_rows.append(["Applied model", applied or "not reported in this output file"])
                else:
                    basis_rows.append(["Source-term option", option])
            if basis_rows:
                tables.append({
                    "section": "Source Term  -  RG 1.183 / NBT",
                    "title": "Source-term option",
                    "df": pd.DataFrame(basis_rows, columns=["Field", "Value"]),
                })

            inv_col = "Core inventory (Ci)"
            rel_col = "Released inventory (Ci)"
            frac_col = "Total release frac"
            if inv_col in st_df.columns:
                main_mask = st_df[inv_col].apply(
                    lambda v: isinstance(v, (int, float, np.integer, np.floating)) and not pd.isna(v)
                )
                main_df = st_df[main_mask].copy()
                main_df = _clean_source_term_display_df(main_df)
                if not main_df.empty:
                    display_cols = [c for c in [
                        "Group", "NBT key", inv_col, "Gap release %", "Early IV %",
                        "Total release %", frac_col, rel_col,
                    ] if c in main_df.columns]
                    fmt = {inv_col: "{:.4e}", rel_col: "{:.4e}", frac_col: "{:.4e}",
                           "Gap release %": "{:.3f}", "Early IV %": "{:.3f}",
                           "Total release %": "{:.3f}"}
                    tables.append({
                        "section": "Source Term  -  RG 1.183 / NBT",
                        "title": "Group inventories and releases",
                        "df": _report_df(main_df[display_cols].reset_index(drop=True), fmt),
                    })

            meta_labels = {
                "Model selected", "Fuel source-term option", "Inventory model",
                "Estimated fissions", "NOTBADTRAD gap release duration (hr)",
                "NOTBADTRAD early-IV release duration (hr)",
                "NOTBADTRAD EAB integration time (hr)", "NOTBADTRAD LPZ integration time (hr)",
                "Severe-event ECR release flag", "Severe-event ECR threshold (%)",
                "Predicted mean oxidizing-rod ECR (%)", "Severe-event ECR active",
            }
            if "Group" in st_df.columns:
                meta_df = st_df[st_df["Group"].isin(meta_labels)].copy()
                if not meta_df.empty:
                    meta_df = _clean_numeric_columns_for_streamlit(
                        meta_df, numeric_cols=[], text_cols=list(meta_df.columns)
                    )
                    meta_cols = [c for c in ["Group", frac_col] if c in meta_df.columns]
                    if meta_cols:
                        tables.append({
                            "section": "Source Term  -  RG 1.183 / NBT",
                            "title": "Source-term basis",
                            "df": _report_df(meta_df[meta_cols].reset_index(drop=True)),
                        })

            if inv_col not in st_df.columns:
                display_cols = [c for c in ["Group", "Gap release %", "Early IV %", "Total release %"] if c in st_df.columns]
                if display_cols:
                    tables.append({
                        "section": "Source Term  -  RG 1.183 / NBT",
                        "title": "Source-term releases",
                        "df": _report_df(_clean_source_term_display_df(st_df[display_cols]).reset_index(drop=True)),
                    })
        except Exception as e:
            tables.append({
                "section": "Source Term  -  RG 1.183 / NBT",
                "title": "Source-term parsing warning",
                "df": pd.DataFrame([[str(e)]], columns=["Warning"]),
            })

    # NOTBADTRAD Dose Screening block
    dose_rows = load_dose(case_name, run_dir=run_dir)
    if dose_rows:
        summary = [r[:5] for r in dose_rows if r and r[0] in ("EAB", "LPZ", "Control Room")]
        if summary:
            sum_df = pd.DataFrame(summary, columns=["Location", "TEDE (rem)", "Limit (rem)", "Margin (rem)", "Result"])
            sum_df = _clean_dose_display_df(sum_df)
            tables.append({
                "section": "NOTBADTRAD Dose Screening",
                "title": "Summary",
                "df": _report_df(sum_df, {"TEDE (rem)":"{:.4e}", "Margin (rem)":"{:.4e}"}, text_cols=["Limit (rem)", "Location", "Result"]),
            })
        grp_hdr = next((i for i, r in enumerate(dose_rows) if r and r[0] == "Group"), None)
        if grp_hdr is not None:
            grp_rows = []
            for r in dose_rows[grp_hdr + 1:]:
                if not any(r): break
                grp_rows.append(r)
            if grp_rows:
                grp_df = pd.DataFrame([r[:6] for r in grp_rows], columns=["Group", "Released (Ci)", "EAB (rem)", "LPZ (rem)", "CR (rem)", "Release frac"])
                grp_df = _clean_dose_display_df(grp_df)
                tables.append({
                    "section": "NOTBADTRAD Dose Screening",
                    "title": "Group contributions",
                    "df": _report_df(grp_df, {"Released (Ci)":"{:.4e}", "EAB (rem)":"{:.4e}", "LPZ (rem)":"{:.4e}", "CR (rem)":"{:.4e}", "Release frac":"{:.4e}"}, text_cols=["Group"]),
                })
        fit_row = next((r for r in dose_rows if r and r[0] == "Power-law fit"), None)
        if fit_row and len(fit_row) > 1 and fit_row[1]:
            fit_err = next((r[1] for r in dose_rows if r and r[0] == "Fit error (%)"), None)
            fit_data = [["Power-law fit", fit_row[1]]]
            if fit_err is not None:
                fit_data.append(["Fit error (%)", fit_err])
            tables.append({
                "section": "NOTBADTRAD Dose Screening",
                "title": "TEDE vs Distance fit",
                "df": pd.DataFrame(fit_data, columns=["Field", "Value"]),
            })
        dist_hdr = next((i for i, r in enumerate(dose_rows) if r and r[0] == "Distance (m)"), None)
        if dist_hdr is not None:
            dt_rows = []
            for r in dose_rows[dist_hdr + 1:]:
                if not any(r): break
                dt_rows.append(r)
            if dt_rows:
                hdr = dose_rows[dist_hdr]
                if len(hdr) >= 4 and "chi" in str(hdr[1]).lower() and "EAB" in str(hdr[2]) and "LPZ" in str(hdr[3]):
                    dt_df = pd.DataFrame([r[:4] for r in dt_rows], columns=["Distance (m)", "χ/Q (s/m³)", "TEDE EAB interval (rem)", "TEDE LPZ interval (rem)"])
                    dt_df = _clean_dose_display_df(dt_df)
                    tables.append({"section": "NOTBADTRAD Dose Screening", "title": "TEDE vs Distance", "df": _report_df(dt_df, {"χ/Q (s/m³)":"{:.4e}", "TEDE EAB interval (rem)":"{:.4e}", "TEDE LPZ interval (rem)":"{:.4e}"})})
                elif len(hdr) >= 3 and "EAB" in str(hdr[1]) and "LPZ" in str(hdr[2]):
                    dt_df = pd.DataFrame([r[:3] for r in dt_rows], columns=["Distance (m)", "TEDE EAB interval (rem)", "TEDE LPZ interval (rem)"])
                    dt_df = _clean_dose_display_df(dt_df)
                    tables.append({"section": "NOTBADTRAD Dose Screening", "title": "TEDE vs Distance", "df": _report_df(dt_df, {"TEDE EAB interval (rem)":"{:.4e}", "TEDE LPZ interval (rem)":"{:.4e}"})})

    # Iodine Spike block
    isp_rows = load_iodine_spike(case_name, run_dir=run_dir)
    if isp_rows:
        def _isp_val(key):
            row = next((r for r in isp_rows if r and r[0] == key), None)
            return row[1] if row and len(row) > 1 else None
        metric_rows = []
        for k in ["Model", "Coolant activity (uCi/g)", "Primary mass (kg)", "Spike inventory (Ci)", "Spike multiplier", "Scram occurred", "PORV opened", "DNB occurred", "N rods DNB/dryout", "Equilibrium spike frac", "Accident spike frac (DNB)"]:
            v = _isp_val(k)
            if v is not None:
                metric_rows.append([k, v])
        if metric_rows:
            tables.append({
                "section": "Iodine Spike — Pre-existing Coolant Activity",
                "title": "Inputs and trigger flags",
                "df": pd.DataFrame(metric_rows, columns=["Field", "Value"]),
            })
        summary = [r[:5] for r in isp_rows if r and r[0] in ("EAB", "LPZ", "Control Room")]
        if summary:
            sum_df = pd.DataFrame(summary, columns=["Location", "TEDE (rem)", "Limit (rem)", "Margin (rem)", "Result"])
            sum_df = _clean_dose_display_df(sum_df)
            tables.append({
                "section": "Iodine Spike — Pre-existing Coolant Activity",
                "title": "Summary",
                "df": _report_df(sum_df, {"TEDE (rem)":"{:.4e}", "Margin (rem)":"{:.4e}"}, text_cols=["Limit (rem)", "Location", "Result"]),
            })
        grp_hdr = next((i for i, r in enumerate(isp_rows) if r and r[0] == "Group"), None)
        if grp_hdr is not None:
            grp_rows = []
            for r in isp_rows[grp_hdr + 1:]:
                if not any(r): break
                grp_rows.append(r)
            if grp_rows:
                grp_df = pd.DataFrame([r[:6] for r in grp_rows], columns=["Group", "Released (Ci)", "EAB (rem)", "LPZ (rem)", "CR (rem)", "Release frac"])
                grp_df = _clean_dose_display_df(grp_df)
                tables.append({
                    "section": "Iodine Spike — Pre-existing Coolant Activity",
                    "title": "Group contributions",
                    "df": _report_df(grp_df, {"Released (Ci)":"{:.4e}", "EAB (rem)":"{:.4e}", "LPZ (rem)":"{:.4e}", "CR (rem)":"{:.4e}", "Release frac":"{:.4e}"}, text_cols=["Group"]),
                })
    return tables

def load_dose(case_name, run_dir=None):
    """Load Dose sheet from output workbook if it exists."""
    for search_dir in ([run_dir] if run_dir is not None else []) + [WORK_DIR]:
        p = Path(search_dir) / f"{case_name}_out.xlsx"
        if p.exists():
            try:
                wb = load_workbook(p, read_only=True)
                if "Dose" not in wb.sheetnames:
                    wb.close()
                    return None
                ws = wb["Dose"]
                rows = [[c.value for c in row] for row in ws.iter_rows()]
                wb.close()
                return rows
            except Exception:
                return None
    return None
    return None


# ── Plot builder ──────────────────────────────────────────────────────────────

def base_layout(title="", height=300):
    return dict(
        title=dict(text=title, font=dict(family="Space Mono", size=12,
                                         color=PLOT_TEXT)),
        plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_PAPER,
        font=dict(family="Inter", size=11, color=PLOT_TEXT),
        xaxis=dict(gridcolor=PLOT_GRID, zerolinecolor=PLOT_GRID,
                   title_font_color=PLOT_TEXT),
        yaxis=dict(gridcolor=PLOT_GRID, zerolinecolor=PLOT_GRID,
                   title_font_color=PLOT_TEXT),
        margin=dict(l=50, r=20, t=40, b=40),
        height=height,
        showlegend=True,
        legend=dict(bgcolor="rgba(0,0,0,0)", font=dict(size=10)),
    )


def make_plots(df, case_name="", use_english=False):
    """Build Plotly figures matching the 16 PNG figures from post_processing."""
    t   = df["Time (s)"].values

    def gcol(col):
        return df[col].values if col in df.columns else np.zeros(len(t))

    def gcol_nan(col):
        if col not in df.columns: return np.full(len(t), np.nan)
        return df[col].replace(0, np.nan).values

    C = ["#0969da","#cf222e","#1a7f37","#9a6700","#6e40c9","#bc4c00"]
    LW = 2.0

    # ── Unit-system scalars ───────────────────────────────────────────────────
    if use_english:
        def K_to_T(K):
            return np.where(np.isfinite(K), (K - 273.15)*9/5 + 32, np.nan)
        T_unit              = "°F"
        P_scale, P_unit     = 0.145038, "psia"
        M_scale, M_unit     = 2.20462,  "lb/s"
        V_scale, V_unit     = 3.28084,  "ft/s"
        L_scale, L_unit     = 3.28084,  "ft"

    else:
        def K_to_T(K):
            return np.where(np.isfinite(K), K - 273.15, np.nan)
        T_unit              = "°C"
        P_scale, P_unit     = 1.0,    "kPa"
        M_scale, M_unit     = 1.0,    "kg/s"
        V_scale, V_unit     = 1.0,    "m/s"
        L_scale, L_unit     = 1.0,    "m"

    def fig_base(title, xlabel="Time [s]", ylabel="", height=300):
        fig = go.Figure()
        fig.update_layout(
            title=dict(text=title, font=dict(family="IBM Plex Mono", size=12, color=PLOT_TEXT)),
            plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_PAPER,
            font=dict(family="IBM Plex Sans", size=11, color=PLOT_TEXT),
            xaxis=dict(title=xlabel, gridcolor=PLOT_GRID, zerolinecolor=PLOT_GRID),
            yaxis=dict(title=ylabel, gridcolor=PLOT_GRID, zerolinecolor=PLOT_GRID),
            margin=dict(l=55, r=20, t=40, b=45),
            height=height,
            legend=dict(bgcolor="rgba(255,255,255,0.8)", bordercolor=PLOT_GRID,
                        borderwidth=1, font=dict(size=10)),
        )
        return fig

    def no_data(fig, msg="No data"):
        fig.add_annotation(text=msg, xref="paper", yref="paper",
                           x=0.5, y=0.5, showarrow=False,
                           font=dict(size=13, color=PLOT_TEXT))

    figs = {}

    # ── Fig 1: RCS Pressure ──────────────────────────────────────────────────
    fig = fig_base("Fig 1  -  RCS Pressure", ylabel=f"Pressure [{P_unit}]")
    fig.add_trace(go.Scatter(x=t, y=gcol("RCS Pressure (kPa)") * P_scale,
                             line=dict(color=C[0], width=LW), name="Pressure"))
    figs["fig1"] = fig

    # ── Fig 2: Equilibrium Quality & Void Fraction ───────────────────────────
    fig = fig_base("Fig 2  -  Equilibrium Quality & Void Fraction",
                   ylabel="Quality / Void Fraction [-]")
    xeq  = gcol("Equilibrium Quality (-)")
    vfrac = gcol("Void Fraction (-)")
    if np.any(xeq != 0) or np.any(vfrac > 0):
        fig.add_trace(go.Scatter(x=t, y=xeq,
                                 line=dict(color=C[0], width=LW), name="X_eq"))
        if np.any(vfrac > 0):
            fig.add_trace(go.Scatter(x=t, y=vfrac,
                                     line=dict(color=C[1], width=LW, dash="dash"),
                                     name="Void Fraction"))
        fig.add_hline(y=0, line=dict(color="grey", width=1, dash="dash"))
        fig.add_hline(y=1, line=dict(color=C[3], width=1, dash="dash"))
    else:
        no_data(fig, "X_eq = 0 throughout (subcooled)")
    figs["fig2"] = fig

    # ── Fig 3: RCS Mass (normalized) ─────────────────────────────────────────
    fig = fig_base("Fig 3  -  RCS Mass (normalized)", ylabel="Mass (normalized)")
    fig.add_trace(go.Scatter(x=t, y=gcol("Total Mass Scaled"),
                             line=dict(color=C[0], width=LW), name="Scaled Mass"))
    figs["fig3"] = fig

    # ── Fig 4: RPV Liquid Level vs Time ──────────────────────────────────────
    fig = fig_base("Fig 4  -  RPV Liquid Level", ylabel=f"Level [{L_unit}]")
    fig.add_trace(go.Scatter(x=t, y=gcol("Vessel Level (m)") * L_scale,
                             line=dict(color=C[0], width=LW), name="Vessel Level"))
    figs["fig4"] = fig

    # ── Fig 5: Break / PORV Flow ──────────────────────────────────────────────
    fig = fig_base("Fig 5  -  Break / PORV Flow Rate", ylabel=f"Mass flow [{M_unit}]")
    fig.add_trace(go.Scatter(x=t, y=gcol("Break Flow (kg/s)") * M_scale,
                             line=dict(color=C[0], width=LW), name="Break flow"))
    porv = gcol("PORV Mass Flow (kg/s)") * M_scale
    if np.any(porv > 0):
        fig.add_trace(go.Scatter(x=t, y=porv,
                                 line=dict(color=C[3], width=LW, dash="dash"),
                                 name="PORV flow"))
    figs["fig5"] = fig

    # ── Fig 6: Accumulator Pressure ───────────────────────────────────────────
    fig = fig_base("Fig 6  -  Accumulator Pressure", ylabel=f"Pressure [{P_unit}]")
    fig.add_trace(go.Scatter(x=t, y=gcol("Accumulator Pressure (kPa)") * P_scale,
                             line=dict(color=C[0], width=LW), name="Acc. Pressure"))
    figs["fig6"] = fig

    # ── Fig 7: Accumulator Temperature ───────────────────────────────────────
    fig = fig_base("Fig 7  -  Accumulator Temperature", ylabel=f"Temperature [{T_unit}]")
    fig.add_trace(go.Scatter(x=t, y=K_to_T(gcol("Accumulator Temperature (K)")),
                             line=dict(color=C[0], width=LW), name="Acc. Temp"))
    figs["fig7"] = fig

    # ── Fig 8: Accumulator Liquid Level ──────────────────────────────────────
    fig = fig_base("Fig 8  -  Accumulator Liquid Level", ylabel=f"Level [{L_unit}]")
    acc_lvl = gcol("Accumulator Level (m)") if "Accumulator Level (m)" in df.columns else None
    if acc_lvl is not None and np.any(acc_lvl > 0):
        fig.add_trace(go.Scatter(x=t, y=acc_lvl * L_scale,
                                 line=dict(color=C[0], width=LW), name="Accumulator Level"))
    else:
        no_data(fig, "No accumulator level data")
    figs["fig8"] = fig

    # ── Fig 9: Accumulator & SI / CVCS Flow Rate ─────────────────────────────
    fig = fig_base("Fig 9  -  Accumulator & SI / CVCS Flow Rate", ylabel=f"Flow [{M_unit}]")
    acc_flow = gcol("Accumulator Flow (kg/s)") * M_scale
    fig.add_trace(go.Scatter(x=t, y=acc_flow,
                             line=dict(color=C[0], width=LW), name="Accumulator"))
    cvcs_mu = gcol("CVCS Makeup (kg/s)")  if "CVCS Makeup (kg/s)"  in df.columns else None
    cvcs_ld = gcol("CVCS Letdown (kg/s)") if "CVCS Letdown (kg/s)" in df.columns else None
    if cvcs_mu is not None and np.any(cvcs_mu > 0):
        fig.add_trace(go.Scatter(x=t, y=cvcs_mu * M_scale,
                                 line=dict(color=C[2], width=LW), name="CVCS Makeup"))
    if cvcs_ld is not None and np.any(cvcs_ld > 0):
        fig.add_trace(go.Scatter(x=t, y=-cvcs_ld * M_scale,
                                 line=dict(color=C[2], width=LW, dash="dot"),
                                 name="CVCS Letdown"))
    hpsi = gcol("HPSI Flow (kg/s)") * M_scale
    if np.any(hpsi > 0):
        fig.add_trace(go.Scatter(x=t, y=hpsi,
                                 line=dict(color=C[3], width=LW), name="HPSI"))
    lpsi = gcol("LPSI Flow (kg/s)") * M_scale
    if np.any(lpsi > 0):
        fig.add_trace(go.Scatter(x=t, y=lpsi,
                                 line=dict(color=C[4], width=LW), name="LPSI"))
    si_total = gcol("SI Pumped Total (kg/s)") * M_scale
    cvcs_net = ((cvcs_mu - cvcs_ld) if (cvcs_mu is not None and cvcs_ld is not None)
                else (cvcs_mu if cvcs_mu is not None else 0))
    total_eccs = acc_flow + (cvcs_net * M_scale if cvcs_net is not None else 0) + si_total
    if np.any(si_total > 0) or (cvcs_net is not None and np.any(cvcs_net != 0)):
        fig.add_trace(go.Scatter(x=t, y=total_eccs,
                                 line=dict(color=C[1], width=LW, dash="dot"),
                                 name="Total ECCS"))
    figs["fig9"] = fig

    # ── Fig 10: Core Power ────────────────────────────────────────────────────
    fig = fig_base("Fig 10  -  Core Power", ylabel="Power [MW]")
    rk = gcol_nan("RK Total Power (MW)")
    if np.any(np.isfinite(rk)):
        fig.add_trace(go.Scatter(x=t, y=rk,
                                 line=dict(color=C[3], width=LW),
                                 name="Fission+Decay Source"))
    fig.add_trace(go.Scatter(x=t, y=gcol("Core Power (MW)"),
                             line=dict(color=C[0], width=LW),
                             name="Convected to Coolant"))
    figs["fig10"] = fig

    # ── Fig 11: RCP Speed ────────────────────────────────────────────────────
    fig = fig_base("Fig 11  -  RCP Coastdown", ylabel="Speed [rpm]")
    ps = gcol("Pump Speed (rpm)")
    if np.any(ps > 0):
        fig.add_trace(go.Scatter(x=t, y=ps,
                                 line=dict(color=C[0], width=LW), name="RCP Speed"))
    else:
        no_data(fig, "No pump data (pump_flag=0)")
    figs["fig11"] = fig

    # ── Fig 12: Coolant Velocity ─────────────────────────────────────────────
    fig = fig_base("Fig 12  -  RCP Coolant Velocity", ylabel=f"Velocity [{V_unit}]")
    pv = gcol("Pump Velocity (m/s)") * V_scale
    if np.any(pv > 0):
        fig.add_trace(go.Scatter(x=t, y=pv,
                                 line=dict(color=C[1], width=LW), name="Coolant Velocity"))
    else:
        no_data(fig, "No pump data (pump_flag=0)")
    figs["fig12"] = fig

    # ── Fig 13: SG Heat Removal ───────────────────────────────────────────────
    fig = fig_base("Fig 13  -  Steam Generator Heat Removal", ylabel="Heat Removal [MW]")
    sg = gcol("SG Heat Removal (MW)")
    if np.any(sg != 0):
        fig.add_trace(go.Scatter(x=t, y=sg,
                                 line=dict(color=C[3], width=LW), name="SG Removal"))
    else:
        no_data(fig, "No SG data (sg_flag=0)")
    figs["fig13"] = fig

    # ── Fig 14: Core Temperatures ─────────────────────────────────────────────
    fig = fig_base("Fig 14  -  Core Temperatures", ylabel=f"Temperature [{T_unit}]")
    tw    = K_to_T(gcol_nan("Clad Surface Temp (K)"))
    tc    = K_to_T(gcol("RCS Temperature (K)"))
    tfuel = K_to_T(gcol_nan("Fuel Avg Temp (K)"))
    thc   = K_to_T(gcol_nan("Hot Pin Clad Temp (K)"))
    thf   = K_to_T(gcol_nan("Hot Pin Fuel Temp (K)"))
    if np.any(np.isfinite(tw)):
        fig.add_trace(go.Scatter(x=t, y=tw,   line=dict(color=C[3], width=LW),          name="Clad surface"))
    fig.add_trace(    go.Scatter(x=t, y=tc,   line=dict(color=C[0], width=LW),          name="Coolant bulk"))
    if np.any(np.isfinite(tfuel)):
        fig.add_trace(go.Scatter(x=t, y=tfuel,line=dict(color=C[2], width=LW, dash="dash"), name="Fuel avg"))
    if np.any(np.isfinite(thc)):
        fig.add_trace(go.Scatter(x=t, y=thc,  line=dict(color=C[1], width=LW, dash="dot"),  name="Hot pin clad"))
    if np.any(np.isfinite(thf)):
        fig.add_trace(go.Scatter(x=t, y=thf,  line=dict(color=C[3], width=LW, dash="dot"),  name="Hot pin fuel"))
    figs["fig14"] = fig

    # ── Fig 15: Core HTC and Power (dual axis) ────────────────────────────────
    # HTC stays in W/m²·K in both unit systems (no simple scalar conversion)
    fig = make_subplots(specs=[[{"secondary_y": True}]])
    htc = gcol("Clad HTC (W/m2-K)")
    if np.any(htc > 0):
        fig.add_trace(go.Scatter(x=t, y=htc,
                                 line=dict(color=C[0], width=LW), name="HTC (left)"),
                      secondary_y=False)
    rk2 = gcol_nan("RK Total Power (MW)")
    if np.any(np.isfinite(rk2)):
        fig.add_trace(go.Scatter(x=t, y=rk2,
                                 line=dict(color=C[3], width=LW, dash="dash"),
                                 name="Heat source MW (right)"), secondary_y=True)
    fig.add_trace(go.Scatter(x=t, y=gcol("Core Power (MW)"),
                             line=dict(color=C[2], width=LW, dash="dash"),
                             name="Convected MW (right)"), secondary_y=True)
    fig.update_layout(
        title=dict(text="Fig 15  -  Core HTC and Power",
                   font=dict(family="IBM Plex Mono", size=12, color=PLOT_TEXT)),
        plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_PAPER,
        font=dict(family="IBM Plex Sans", size=11, color=PLOT_TEXT),
        margin=dict(l=55, r=65, t=40, b=45), height=300,
        legend=dict(bgcolor="rgba(255,255,255,0.8)"),
    )
    fig.update_yaxes(title_text="HTC [W/m²·K]", secondary_y=False,
                     gridcolor=PLOT_GRID, title_font_color=C[0])
    fig.update_yaxes(title_text="Power [MW]",   secondary_y=True,
                     gridcolor=PLOT_GRID, title_font_color=C[3])
    fig.update_xaxes(title_text="Time [s]", gridcolor=PLOT_GRID)
    figs["fig15"] = fig

    # ── Fig 16: DNBR ─────────────────────────────────────────────────────────
    fig = fig_base("Fig 16  -  DNBR (Biasi/Zuber/Bowring)", ylabel="DNBR [-]")
    dn = gcol_nan("DNBR")
    if np.any(np.isfinite(dn)):
        fig.add_trace(go.Scatter(x=t, y=dn,
                                 line=dict(color=C[0], width=LW), name="DNBR"))
        fig.add_hline(y=1.0, line=dict(color="#cf222e", width=2, dash="dash"),
                      annotation_text="CHF limit (1.0)",
                      annotation_font_color="#cf222e")
        fig.add_hline(y=1.3, line=dict(color="#9a6700", width=1.5, dash="dot"),
                      annotation_text="Design limit (1.3)",
                      annotation_font_color="#9a6700")
    else:
        no_data(fig, "No DNBR data (dnbr_flag=0)")
    figs["fig16"] = fig

    # ── Fig 17: Reactivity ────────────────────────────────────────────────────
    fig = fig_base("Fig 17  -  Reactivity", ylabel="Reactivity [pcm]")
    rscram= gcol("Reactivity scram (pcm)")
    rext  = gcol("Reactivity ext (pcm)")
    rboron= gcol("Reactivity Boron (pcm)")
    rdop  = gcol("Reactivity Doppler (pcm)")
    rmod  = gcol("Reactivity Moderator (pcm)")
    rnet  = gcol("Reactivity net (pcm)")
    if np.any(rscram != 0) or np.any(rnet != 0) or np.any(rext != 0) or np.any(rboron != 0):
        if np.any(rscram != 0):
            fig.add_trace(go.Scatter(x=t, y=rscram, line=dict(color=C[1], width=LW),
                                     name="ρ_scram"))
        if np.any(rext != 0):
            fig.add_trace(go.Scatter(x=t, y=rext, line=dict(color="grey", width=LW,
                                                              dash="dash"),
                                     name="ρ_ext (rod withdrawal)"))
        if np.any(rboron != 0):
            fig.add_trace(go.Scatter(x=t, y=rboron, line=dict(color="purple", width=LW,
                                                               dash="dashdot"),
                                     name="ρ_boron (SLCS)"))
        if np.any(rdop != 0):
            fig.add_trace(go.Scatter(x=t, y=rdop, line=dict(color=C[3], width=LW,
                                                             dash="dash"), name="ρ_Doppler"))
        if np.any(rmod != 0):
            fig.add_trace(go.Scatter(x=t, y=rmod, line=dict(color=C[0], width=LW,
                                                             dash="dot"), name="ρ_moderator"))
        if np.any(rnet != 0):
            fig.add_trace(go.Scatter(x=t, y=rnet, line=dict(color="black", width=LW+0.5),
                                     name="ρ_net"))
        fig.add_hline(y=0, line=dict(color="grey", width=1))
    else:
        no_data(fig, "No reactivity data")
    figs["fig17"] = fig

    # ── Fig pzr: Pressurizer Level (always shown when data present) ──────────
    pzr_m    = gcol_nan("Pressurizer Level (m)")
    pzr_norm = gcol_nan("Pressurizer Level (norm)")
    if np.any(np.isfinite(pzr_m)):
        fig = make_subplots(specs=[[{"secondary_y": True}]])
        fig.add_trace(go.Scatter(x=t, y=pzr_m * L_scale,
                                 line=dict(color=C[0], width=LW),
                                 name=f"Level ({L_unit})"),
                      secondary_y=False)
        fig.add_trace(go.Scatter(x=t, y=pzr_norm,
                                 line=dict(color=C[3], width=LW, dash="dash"),
                                 name="Level (norm)"),
                      secondary_y=True)
        fig.update_layout(
            title=dict(text="Pressurizer Level",
                       font=dict(family="IBM Plex Mono", size=12, color=PLOT_TEXT)),
            plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_PAPER,
            font=dict(family="IBM Plex Sans", size=11, color=PLOT_TEXT),
            margin=dict(l=55, r=65, t=40, b=45), height=300,
            legend=dict(bgcolor="rgba(255,255,255,0.8)"),
        )
        fig.update_yaxes(title_text=f"Level [{L_unit}]", secondary_y=False,
                         gridcolor=PLOT_GRID, title_font_color=C[0])
        fig.update_yaxes(title_text="Level (norm)", secondary_y=True,
                         gridcolor=PLOT_GRID, title_font_color=C[3])
        fig.update_xaxes(title_text="Time [s]", gridcolor=PLOT_GRID)
        figs["fig_pzr"] = fig

    # ── Zircaloy oxidation / hydrogen generation plots ──────────────────────
    if ("Zr Oxidation Hot Pin ECR (%)" in df.columns and
            np.nanmax(gcol("Zr Oxidation Hot Pin ECR (%)")) > 0):
        ecr_hot = gcol("Zr Oxidation Hot Pin ECR (%)")
        ecr_oxid = gcol("Zr Oxidation Mean Oxidizing Rod ECR (%)")
        fig = fig_base("Zircaloy Oxidation / ECR", ylabel="ECR [%]")
        fig.add_trace(go.Scatter(x=t, y=ecr_hot, line=dict(color=C[3], width=LW),
                                 name="Hot-pin ECR"))
        if np.nanmax(ecr_oxid) > 0:
            fig.add_trace(go.Scatter(x=t, y=ecr_oxid, line=dict(color=C[0], width=LW, dash="dash"),
                                     name="Mean oxidizing rod ECR"))
        fig.add_hline(y=17.0, line_dash="dot", line_color="red",
                      annotation_text="17% ECR", annotation_position="top right")
        figs["fig18_zr_ecr"] = fig

    if ("H2 Generated (kg)" in df.columns and
            np.nanmax(gcol("H2 Generated (kg)")) > 0):
        h2 = gcol("H2 Generated (kg)")
        if "H2 Full Core Cladding Reaction (kg)" in df.columns:
            _h2_full_for_cap = gcol("H2 Full Core Cladding Reaction (kg)")
            if np.nanmax(_h2_full_for_cap) > 0:
                h2 = np.minimum(h2, _h2_full_for_cap)
        n_oxid = gcol("Zr Oxidizing Rods (est.)")
        fig = make_subplots(specs=[[{"secondary_y": True}]])
        fig.add_trace(go.Scatter(x=t, y=h2, line=dict(color=C[1], width=LW),
                                 name="H₂ generated"),
                      secondary_y=False)
        if "H2 Full Core Cladding Reaction (kg)" in df.columns:
            h2_full = gcol("H2 Full Core Cladding Reaction (kg)")
            if np.nanmax(h2_full) > 0:
                fig.add_trace(go.Scatter(x=t, y=h2_full, line=dict(color=C[3], width=LW, dash="dot"),
                                         name="Full-core clad reaction H₂"),
                              secondary_y=False)
        if np.nanmax(n_oxid) > 0:
            fig.add_trace(go.Scatter(x=t, y=n_oxid, line=dict(color=C[0], width=LW, dash="dash"),
                                     name="Oxidizing rods"),
                          secondary_y=True)
        fig.update_layout(
            title=dict(text="Hydrogen Generation",
                       font=dict(family="IBM Plex Mono", size=12, color=PLOT_TEXT)),
            plot_bgcolor=PLOT_BG, paper_bgcolor=PLOT_PAPER,
            font=dict(family="IBM Plex Sans", size=11, color=PLOT_TEXT),
            margin=dict(l=55, r=65, t=40, b=45), height=300,
            legend=dict(bgcolor="rgba(255,255,255,0.8)"),
        )
        fig.update_yaxes(title_text="H₂ generated [kg]", secondary_y=False,
                         gridcolor=PLOT_GRID, title_font_color=C[1])
        fig.update_yaxes(title_text="Oxidizing rods [count]", secondary_y=True,
                         gridcolor=PLOT_GRID, title_font_color=C[0])
        fig.update_xaxes(title_text="Time [s]", gridcolor=PLOT_GRID)
        figs["fig19_h2_generation"] = fig


    # Apply overlay legend styling to all figures
    for f in figs.values():
        try:
            f.update_layout(
                legend=dict(
                    x=0.98, y=0.98, xanchor="right", yanchor="top",
                    bgcolor="rgba(0,0,0,0)", borderwidth=0, font=dict(size=10),
                )
            )
        except Exception:
            pass

    return figs


# ── Scalar summary ────────────────────────────────────────────────────────────

def scalar_summary(df, use_english=False, init_power_mw=None, diag_df=None):
    """
    Return list of (display_value, label, severity_class) metric tiles.

    init_power_mw  - rated power from the input file, used as the baseline
                     peak power for cases that immediately trip (e.g. LOCA).
    diag_df        - optional <case>_diag.csv data used for SG flow-capacity
                     status on the Peak Power tile.
    """
    metrics = []

    # ── MDNBR ────────────────────────────────────────────────────────────────
    if "DNBR" in df.columns:
        dn = df["DNBR"].replace(0, np.nan)
        mdnbr = dn.min()
        if not np.isnan(mdnbr):
            t_min = df["Time (s)"].iloc[dn.idxmin()]
            cls = "danger" if mdnbr < 1.0 else "warn" if mdnbr < 1.5 else "ok"
            metrics.append((f"{mdnbr:.3f}", f"MDNBR  (t={t_min:.1f}s)", cls))

    # ── Peak Clad Temperature ─────────────────────────────────────────────────
    # ── Peak Clad Temperature (avg + hot pin combined) ────────────────────
    _has_avg = "Clad Surface Temp (K)" in df.columns
    _has_hot = "Hot Pin Clad Temp (K)" in df.columns
    if _has_avg or _has_hot:
        tw_avg_K = df["Clad Surface Temp (K)"].replace(0, np.nan).max() if _has_avg else np.nan
        tw_hot_K = df["Hot Pin Clad Temp (K)"].replace(0, np.nan).max() if _has_hot else np.nan
        _ref_K   = tw_hot_K if not np.isnan(tw_hot_K) else tw_avg_K
        if not np.isnan(_ref_K):
            cls = ("danger" if _ref_K - 273.15 > 1500
                   else "warn" if _ref_K - 273.15 > 1204 else "ok")
            if use_english:
                def _to_f(k): return f"{(k-273.15)*9/5+32:.0f}°F"
                avg_str = _to_f(tw_avg_K) if not np.isnan(tw_avg_K) else "—"
                hot_str = _to_f(tw_hot_K) if not np.isnan(tw_hot_K) else "—"
            else:
                def _to_c(k): return f"{k-273.15:.0f}°C"
                avg_str = _to_c(tw_avg_K) if not np.isnan(tw_avg_K) else "—"
                hot_str = _to_c(tw_hot_K) if not np.isnan(tw_hot_K) else "—"
            if _has_avg and _has_hot and not np.isnan(tw_hot_K):
                val_str = f"Avg {avg_str} / Hot {hot_str}"
            elif _has_hot and not np.isnan(tw_hot_K):
                val_str = hot_str
            else:
                val_str = avg_str
            metrics.append((val_str, "Peak Clad Temp (Avg / Hot Pin)", cls, True))

    # ── Peak Pressure ─────────────────────────────────────────────────────────
    if "RCS Pressure (kPa)" in df.columns:
        pmax_kPa = df["RCS Pressure (kPa)"].max()
        if use_english:
            metrics.append((f"{pmax_kPa * 0.145038:.1f}", "Peak Pressure (psia)", "ok"))
        else:
            metrics.append((f"{pmax_kPa:.0f}", "Peak Pressure (kPa)", "ok"))

    # ── Peak Power / SG primary-flow capacity status ─────────────────────────
    # FIX: for cases that trip at t=0 (e.g. LOCA), the engine may never write
    # the rated power into RK Total Power (MW) — the column can start at
    # decay-heat level.  Use the rated power from the input file as a baseline,
    # then compare the displayed peak power to the coincident SG primary-flow
    # capacity limit when diagnostic data are available.
    if "RK Total Power (MW)" in df.columns:
        rk_series = pd.to_numeric(df["RK Total Power (MW)"], errors="coerce")
        rk_max = float(rk_series.max()) if rk_series.notna().any() else 0.0
        init_power = float(init_power_mw) if init_power_mw else 0.0
        pmax = max(init_power, rk_max if rk_max > 0 else 0.0)

        if pmax > 0:
            cap_at_peak = np.nan
            t_peak = None
            if diag_df is not None and "SG Flow-Capacity Limit (MW)" in diag_df.columns:
                cap_series = pd.to_numeric(
                    diag_df["SG Flow-Capacity Limit (MW)"], errors="coerce"
                ).replace([np.inf, -np.inf], np.nan)

                if cap_series.notna().any():
                    # Choose the SG capacity limit at the time represented by the
                    # displayed peak power.  If pmax is the input baseline, use
                    # the first diagnostic row; otherwise use the row where RK
                    # power peaks.
                    if init_power >= rk_max and len(cap_series) > 0:
                        idx_peak = cap_series.first_valid_index()
                    else:
                        idx_peak = rk_series.idxmax() if rk_series.notna().any() else cap_series.first_valid_index()
                        if idx_peak not in cap_series.index or pd.isna(cap_series.loc[idx_peak]):
                            idx_peak = cap_series.first_valid_index()

                    if idx_peak is not None and idx_peak in cap_series.index:
                        cap_at_peak = float(cap_series.loc[idx_peak])
                        if "Time (s)" in diag_df.columns:
                            try:
                                t_peak = float(pd.to_numeric(diag_df["Time (s)"], errors="coerce").loc[idx_peak])
                            except Exception:
                                t_peak = None

            if np.isfinite(cap_at_peak) and cap_at_peak > 0.0:
                if pmax < cap_at_peak:
                    cls = "ok"
                elif pmax < 1.2 * cap_at_peak:
                    cls = "warn"
                else:
                    cls = "danger"
                val = f"{pmax:.0f} / {cap_at_peak:.0f} MW"
                lbl = "Peak Power / SG Flow-Cap Limit"
                if t_peak is not None and np.isfinite(t_peak):
                    lbl += f"  (t≈{t_peak:.1f}s)"
                metrics.append((val, lbl, cls, True))
            else:
                metrics.append((f"{pmax:.0f} MW", "Peak Power", "ok"))

    # ── Rod failure counts  -  always shown (zero is a meaningful result) ─────
    # ── Rod failure counts ─────────────────────────────────────────────────────────────
    # DNB vs Dryout: if a break is present (LOCA) use Dryout label, else DNB
    _has_break = ("Break Flow (kg/s)" in df.columns
                  and df["Break Flow (kg/s)"].max() > 0.1)
    # ── Rod failure counts ─────────────────────────────────────────────
    # DNB vs Dryout: LOCA cases show Dryout, non-LOCA show DNB
    _has_break = ("Break Flow (kg/s)" in df.columns
                  and df["Break Flow (kg/s)"].max() > 0.1)
    _dnb_label = "Dryout" if _has_break else "DNB"
    # DNB/Dryout: reported as impacted rods only (not a gap release trigger)
    # Gap release: only T(clad) > 800°C; Early IV: T(clad) > 1204°C
    _fail_cols = [
        ("Rod Failures DNB (est.)",     f"Rods with {_dnb_label}",              False),
        ("Rod Failures PCT (est.)",     "T(clad) ≥1204°C  (early in-vessel)",   True),
        ("Rod Failures Gap (est.)",     "T(clad) ≥800°C  (gap release)",        True),
        ("Rod Failures EarlyIV (est.)", "T(clad) ≥1204°C  (early in-vessel)",   True),
    ]
    _seen_eiv = False
    for col, label, small in _fail_cols:
        # Avoid showing early IV twice (PCT and EarlyIV map to same thing)
        if label.startswith("T(clad) ≥1204"):
            if _seen_eiv:
                continue
            _seen_eiv = True
        if col in df.columns:
            n = int(df[col].max())
            cls = "danger" if n > 1000 else "warn" if n > 0 else "ok"
            metrics.append((f"{n:,}", label, cls, small))

    # ── Zircaloy oxidation / hydrogen generation ────────────────────────────
    if "Zr Oxidation Hot Pin ECR (%)" in df.columns:
        ecr_max = df["Zr Oxidation Hot Pin ECR (%)"].replace(0, np.nan).max()
        if not np.isnan(ecr_max) and ecr_max > 0:
            cls = "danger" if ecr_max >= 17.0 else "warn" if ecr_max >= 1.0 else "ok"
            metrics.append((f"{ecr_max:.2f}%", "Max Hot-Pin ECR", cls))

    if "H2 Generated (kg)" in df.columns:
        h2_max = df["H2 Generated (kg)"].replace(0, np.nan).max()
        if not np.isnan(h2_max) and h2_max > 0:
            if "H2 Full Core Cladding Reaction (kg)" in df.columns:
                h2_full = df["H2 Full Core Cladding Reaction (kg)"].replace(0, np.nan).max()
            else:
                h2_full = np.nan
            if not np.isnan(h2_full) and h2_full > 0:
                h2_max = min(h2_max, h2_full)
            cls = "danger" if h2_max >= 100.0 else "warn" if h2_max >= 10.0 else "ok"
            if not np.isnan(h2_full) and h2_full > 0:
                metrics.append((f"{h2_max:.2f} kg / {h2_full:.0f} kg",
                                "H₂ Generated / Full-Clad Rxn", cls, True))
            else:
                metrics.append((f"{h2_max:.2f} kg", "Total H₂ Generated", cls))

    return metrics


# ── Sidebar ───────────────────────────────────────────────────────────────────

with st.sidebar:
    # ── Return to FLARE home ──────────────────────────────────────────────────
    st.markdown("""
        <style>
        [data-testid="stSidebar"] button[kind="secondary"] {
            background: transparent !important;
            border: 1px solid #e8530a !important;
            border-radius: 4px !important;
            color: #f97316 !important;
            font-size: 0.82rem !important;
            letter-spacing: 0.08em !important;
            font-weight: 700 !important;
        }
        [data-testid="stSidebar"] button[kind="secondary"]:hover {
            background: rgba(232,83,10,0.18) !important;
            box-shadow: 0 0 14px rgba(232,83,10,0.45) !important;
            color: #ffffff !important;
        }
        </style>""", unsafe_allow_html=True)
    if st.button("\U0001f525  FLARE Home", key="home_btn", width="stretch"):
        st.session_state.page = "home"
        st.query_params.clear()
        st.rerun()
    st.divider()
    # ─────────────────────────────────────────────────────────────────────────
    st.markdown('<div class="app-title">⚛ flare</div>', unsafe_allow_html=True)
    st.markdown('<div class="app-sub">Reactor Safety Analysis</div>',
                unsafe_allow_html=True)

    _case_entries = discover_input_cases()
    if not _case_entries:
        st.error("No *_in.xlsx files found in FLARE subfolders. Root-level input files are intentionally ignored.")
        st.stop()

    _case_labels = [e["label"] for e in _case_entries]
    _selected_label = st.selectbox("Case", _case_labels,
                             help=(
                                 "Input decks are discovered recursively in subfolders below the FLARE root. "
                                 "Root-level *_in.xlsx files and generated sim_/risk_/ua_ folders are ignored. "
                                 "Switching cases clears the previous run output and console log."
                             ))
    _selected_entry = next(e for e in _case_entries if e["label"] == _selected_label)
    selected = _selected_entry["case"]
    selected_input_path = _selected_entry["path"]
    st.caption(f"Input file: `{selected_input_path.relative_to(WORK_DIR)}`")

    # ── Reset Run tab whenever a different case is selected ───────────────────
    # Issue #2: changing the case must clear the Run state and console output.
    if selected != st.session_state.get("_sidebar_case"):
        st.session_state._sidebar_case  = selected
        st.session_state.run_status     = "idle"
        st.session_state.console_log    = ""
        st.session_state.last_case      = None
        st.session_state.run_dir        = None
        st.session_state.run_init_power = None

    st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)

    # ── Previous run selector ─────────────────────────────────────────────────
    # Discover completed sim_* run folders across ALL cases.  This is important
    # for Run All Cases: the batch worker writes normal per-case folders
    # (sim_<Case>_<timestamp>), so a case-specific filter would show only the
    # currently selected case and make it look as if the batch produced just one
    # loadable run.
    _prev_run_entries = []  # list of (label, run_dir, case_name, mtime)
    for _d in WORK_DIR.iterdir():
        if not (_d.is_dir() and _d.name.startswith("sim_")):
            continue
        # Prefer *_out.csv; fall back to *_out.xlsx if the csv is missing.
        # Deduplicate by case name so a run with both files appears only once.
        _seen_cases = set()
        _csv_cases = {f.stem.replace("_out", "") for f in _d.glob("*_out.csv")}
        _xlsx_cases = {f.stem.replace("_out", "") for f in _d.glob("*_out.xlsx")}
        for _case_name in sorted(_csv_cases | _xlsx_cases):
            if not _case_name or _case_name.startswith(".~"):
                continue
            if _case_name in _seen_cases:
                continue
            _seen_cases.add(_case_name)
            _label = f"{_case_name}  /  {_d.name}"
            _prev_run_entries.append((_label, _d, _case_name, _d.stat().st_mtime))

    _prev_run_entries.sort(key=lambda x: x[3], reverse=True)

    if _prev_run_entries:
        _run_labels = ["— select a previous run —"] + [e[0] for e in _prev_run_entries]
        _sel_prev = st.selectbox(
            "Load previous run",
            _run_labels,
            key="prev_run_sel",
            help=(
                "Load results from any previous simulator run. Run All Cases "
                "outputs appear here as ordinary per-case sim_<Case>_<time> folders."
            ),
        )
        if _sel_prev != "— select a previous run —":
            _prev_label, _prev_dir, _prev_case, _prev_mtime = next(
                e for e in _prev_run_entries if e[0] == _sel_prev
            )
            if st.button("Load", key="load_prev_run"):
                st.session_state.last_case      = _prev_case
                st.session_state.run_dir        = _prev_dir
                st.session_state.run_status     = "done"
                st.session_state.console_log    = f"Loaded {_prev_case} from {_prev_dir.name}"
                st.rerun()
        st.caption(f"Found {len(_prev_run_entries)} loadable previous run(s).")
    else:
        st.caption("No previous simulator runs found.")

    st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)

    rows   = load_command_block(selected, selected_input_path)
    params = parse_params(rows)

    KEY_PARAMS = [
        ("endtime",         "End time [s]",
         "Simulation end time in seconds. The engine uses a fine grid (0.01 s steps) "
         "near t=0 and 1 s steps thereafter, so runtime scales roughly linearly."),
        ("F_r",             "Radial peaking F_r",
         "Radial hot-channel factor: ratio of peak-to-average pin power across the core. "
         "Multiplied by F_z to give the combined hot-pin heat flux used in DNBR and PCT."),
        ("F_z",             "Axial peaking F_z",
         "Axial peaking factor: ratio of peak-to-average axial heat flux (chopped cosine). "
         "Typical fresh-core value ~1.55. Combined with F_r for hot-pin calculations."),
        ("h_gap",           "Gap conductance [W/m²K]",
         "Fuel-clad gap conductance. Controls the temperature drop across the as-fabricated "
         "He-filled gap. Lower values raise fuel centreline temperature; "
         "typical BOL range 3000–9000 W/m²K."),
        ("trip_delay",      "Trip delay [s]",
         "Signal processing and relay delay applied to every reactor protection system trip. "
         "The scram fires this many seconds after the setpoint is first exceeded. "
         "Typical hard-wired PWR RPS: 1.5 s."),
        ("trip_power_frac", "High-power trip frac",
         "High-flux scram setpoint as a fraction of rated power "
         "(e.g. 1.10 = trip at 110%). Applies to both the RPS evaluation "
         "and the point-kinetics high-power check in RIA cases."),
        ("trip_P_lo_kPa",   "Low-P trip [kPa]",
         "Low-pressure scram setpoint [kPa]. The reactor trips when RCS pressure "
         "falls below this value. Set to 0 to disable. "
         "Default is 90% of the initial RCS pressure."),
        ("total_power",     "Rated power [MW]",
         "Rated thermal power of the reactor [MW]. Used to normalise decay heat, "
         "set the initial heat flux for core temperatures, and scale rod-failure counts."),
        ("pressure_kPa",    "Initial pressure [kPa]",
         "Initial RCS pressure at t=0 [kPa]. "
         "Standard PWR full-power operating pressure ~15 500 kPa (2250 psia)."),
        ("temp_core_exit",  "Core-exit temp [°C]",
         "Initial bulk coolant temperature at the core exit [°C]. "
         "Sets the initial enthalpy and the starting point for "
         "moderator reactivity feedback in RIA cases."),
        ("alpha_D_pcm",     "Doppler coeff [pcm/°C]",
         "Doppler (fuel temperature) reactivity coefficient [pcm/°C]. "
         "Negative value provides prompt negative feedback during power excursions. "
         "Typical PWR value: −2 to −4 pcm/°C."),
        ("alpha_M_pcm",     "MTC [pcm/°C]",
         "Moderator temperature coefficient [pcm/°C]. "
         "Must be negative at power for inherent stability. "
         "Typical PWR full-power value: −20 to −40 pcm/°C."),
        ("k_sigma",         "k_sigma (failure dist.)",
         "Number of standard deviations between the core-average pin power (f=1) "
         "and the hot pin (f=F_r) in the Gaussian radial power distribution "
         "used to estimate failed-rod counts. Default 3.0 places F_r at the "
         "99.87th percentile."),
    ]

    overrides = {}
    st.caption(
        "Override any numeric parameter for this run only — "
        "the base input file is never modified. "
        "Only parameters present in the input file are shown."
    )
    with st.expander("Edit parameters", expanded=False):
        for key, label, tip in KEY_PARAMS:
            if key in params:
                # Step size = place value of lowest significant digit
                new_val = st.number_input(
                    label,
                    value=float(params[key]),
                    step=smart_step(params[key]),
                    format="%.6g",
                    key=f"ov_{selected}_{key}",
                    help=tip,
                )
                if new_val != params[key]:
                    overrides[key] = new_val

    st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)
    _fast_mode = st.checkbox(
        "Fast mode (no figures)",
        value=False,
        key="fast_mode",
        help=(
            "Skip figure generation and PDF assembly. "
            "Only CSV and XLSX output are written. "
            "Recommended for Uncertainty Analysis and Risk batch runs "
            "where figures are not needed and speed matters."
        ),
    )
    run_btn = st.button("▶  Run Simulation", type="primary", width="stretch")

    st.markdown("### Batch")
    final_report_flag = st.checkbox(
        "Run All - Final Report",
        value=False,
        key="sim_all_final_report",
        help=(
            "When enabled, Run All Cases forces figures ON, generates an AI event narrative "
            "and tables for each case, and compiles all narratives and plots into one final PDF report."
        ),
    )
    _final_report_detail = 0.5
    if final_report_flag:
        _final_report_detail = st.slider(
            "Final Report AI detail level",
            min_value=0.0,
            max_value=1.0,
            value=0.5,
            step=0.05,
            key="sim_all_final_report_detail",
            help=(
                "Controls the detail level for the AI event narratives generated for every case in the final report. "
                "This works like the AI Event Narrative detail slider in the Results panel: lower values produce brief summaries; "
                "higher values produce more detailed technical narratives."
            ),
        )
        if _final_report_detail < 0.25:
            st.caption("Final Report narrative detail: **Brief**")
        elif _final_report_detail < 0.75:
            st.caption("Final Report narrative detail: **Standard**")
        else:
            st.caption("Final Report narrative detail: **Detailed**")
    if final_report_flag and _fast_mode:
        st.caption("Final Report mode overrides Fast mode so figures and plots are generated.")
    run_all_btn = st.button(
        "▶  Run All Cases",
        width="stretch",
        help=(
            "Run every *_in.xlsx case in a durable background worker. "
            "Progress is written to disk so the run can survive browser disconnects."
        ),
    )
    _sim_all_status_sidebar = _sim_all_load_status()
    if _sim_all_status_sidebar and _sim_all_status_sidebar.get("status") in {"starting", "running"}:
        st.caption(
            f"Batch active: {_sim_all_status_sidebar.get('completed_runs', 0)}/"
            f"{_sim_all_status_sidebar.get('total_runs', '?')} complete"
        )

    st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)
    st.caption("flare_sim.py")
    st.caption(f"Working dir: `{WORK_DIR}`")


# ── Main panel ─────────────────────────────────────────────────────────────────

st.markdown(f"## {selected}")

tab_run, tab_results, tab_params = st.tabs(
    ["▶  Run", "📊  Results", "⚙  Input Parameters"]
)

# ── Tab: Input Parameters ─────────────────────────────────────────────────────
with tab_params:
    st.markdown('<div class="card-title">Command block</div>',
                unsafe_allow_html=True)
    if rows:
        lines = "\n".join(t for _, t in rows)
        st.code(lines, language="python")
    else:
        st.info("No parameters loaded.")

# ── Tab: Run ──────────────────────────────────────────────────────────────────
with tab_run:
    # Initialise session-state keys that may not exist yet
    if "console_log"    not in st.session_state: st.session_state.console_log    = ""
    if "run_status"     not in st.session_state: st.session_state.run_status     = "idle"
    if "last_case"      not in st.session_state: st.session_state.last_case      = None
    if "run_dir"        not in st.session_state: st.session_state.run_dir        = None
    if "run_init_power" not in st.session_state: st.session_state.run_init_power = None

    status = st.session_state.run_status

    col1, col2 = st.columns([3, 1])
    with col1:
        if status == "running":
            st.markdown('<span class="badge-run">RUNNING</span>',
                        unsafe_allow_html=True)
        elif status == "done":
            st.markdown('<span class="badge-ok">COMPLETE</span>',
                        unsafe_allow_html=True)
        elif status == "error":
            st.markdown('<span class="badge-fail">FAILED</span>',
                        unsafe_allow_html=True)
        else:
            st.markdown('<span class="badge-warn">IDLE</span>',
                        unsafe_allow_html=True)

    # Durable Run-All-Cases batch mode. This is separate from the single-case
    # Run Simulation path below so existing single-case behavior is unchanged.
    if run_all_btn:
        _active_status = _sim_all_load_status()
        if _active_status and _active_status.get("status") in {"starting", "running"}:
            st.warning("A Run All Cases batch is already active. Abort or wait for it to finish before starting another.")
        else:
            _tag = datetime.now().strftime("%Y%m%d_%H%M%S")
            _batch_dir = WORK_DIR / f"sim_all_{_tag}"
            if _sim_all_launch_worker(_batch_dir, _case_entries, _fast_mode, final_report=final_report_flag, final_report_detail=_final_report_detail):
                st.session_state.sim_all_run_dir = _batch_dir
                st.rerun()

    _sim_all_status = _sim_all_load_status()
    if _sim_all_status and _sim_all_status.get("status") in {"starting", "running", "complete", "aborted", "failed", "error"}:
        _sim_all_render_progress(_sim_all_status)
        st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)

    if run_btn and status != "running":
        st.session_state.console_log = ""
        st.session_state.run_status  = "running"
        st.session_state.sim_time_s  = 0.0

        # Build run case (with overrides if any — minimum_output handled via CLI flag)
        run_case = build_override_xlsx(selected, overrides, selected_input_path) if overrides else selected
        st.session_state.last_case = run_case

        # Issue #4 fix: save rated power so scalar_summary can use it as the
        # peak-power baseline for cases that trip immediately.
        st.session_state.run_init_power = params.get("total_power")

        # Issue #1: create a unique output subfolder for this run
        _tag     = datetime.now().strftime("%Y%m%d_%H%M%S")
        _run_dir = WORK_DIR / f"sim_{selected}_{_tag}"
        _run_dir.mkdir(exist_ok=True)
        st.session_state.run_dir = _run_dir

        # ── Prepare run folder and launch simulation there ────────────────────
        # The simulation now runs directly inside the unique sim_<case>_<timestamp>
        # folder. This avoids moving output files after the subprocess exits and
        # avoids WinError 32 file-lock collisions with Excel, Streamlit, OneDrive,
        # or antivirus indexing.
        _run_input_src = (WORK_DIR / f"{run_case}_in.xlsx") if overrides else selected_input_path
        _run_input_dst = _run_dir / f"{run_case}_in.xlsx"

        try:
            shutil.copy2(str(_run_input_src), str(_run_input_dst))
        except PermissionError as _e:
            st.session_state.run_status = "error"
            st.session_state.console_log = (
                f"PermissionError while copying input workbook to run folder:\n"
                f"  Source: {_run_input_src}\n"
                f"  Target: {_run_input_dst}\n\n"
                f"{_e}\n\n"
                "Close the workbook in Excel or save/close it, then rerun."
            )
            st.rerun()

        # ── Live-streaming subprocess ─────────────────────────────────────────
        _env = os.environ.copy()
        _env["PYTHONUTF8"] = "1"
        _status_slot = col1.empty()
        _status_slot.info("⚙️  PRE-PROCESSING…")
        _sim_cmd = [sys.executable, str(WORK_DIR / "flare_sim.py"), run_case]
        if _fast_mode:
            _sim_cmd.append("--no-figures")
        _proc = subprocess.Popen(
            _sim_cmd,
            stdout=subprocess.PIPE,
            stderr=subprocess.STDOUT,
            text=True, encoding="utf-8",
            cwd=str(_run_dir), env=_env,
        )
        _lines = []
        _sim_started = False
        for _line in _proc.stdout:
            _lines.append(_line)
            if _line.startswith("SIMTIME "):
                _sim_started = True
                try:
                    _t = float(_line.split()[1])
                    _status_slot.info(f"🔄  RUNNING  t = {_t:.0f} s")
                except ValueError:
                    pass
            elif _line.strip() == "POST-PROCESSING":
                _status_slot.info("📊  POST-PROCESSING…")
        _proc.wait()
        combined = "".join(_lines)
        st.session_state.console_log = combined

        if "Traceback" in combined or "Error" in combined.split("\n")[0]:
            st.session_state.run_status = "error"
        else:
            st.session_state.run_status = "done"

        # Outputs are already written directly into _run_dir because the
        # subprocess cwd is _run_dir. Do not move files after the run; moving
        # recently written Excel files is a common source of WinError 32 on
        # Windows/OneDrive systems.

        # Clean up temp override input file in the FLARE root only.  The run
        # folder keeps its own archived copy for traceability.
        if overrides:
            tmp_xlsx = WORK_DIR / f"{run_case}_in.xlsx"
            try:
                tmp_xlsx.unlink()
            except Exception:
                pass

        st.rerun()

    # Console output  -  parse key events for summary
    if st.session_state.console_log:
        log = st.session_state.console_log
        summary_items = []
        import re as _re

        m_exec  = _re.search(r"Total execution time:\s*([\d.]+)\s*s", log)
        if m_exec:  summary_items.append(f"⏱ {m_exec.group(1)} s")

        m_scram = _re.search(r"Reactor scram at t=([\d.]+)", log)
        if m_scram: summary_items.append(f"⚡ Scram t={m_scram.group(1)} s")

        m_dnbr  = _re.search(r"DNBR:\s*min\s*=\s*([\d.]+)\s*at t=([\d.]+)", log)
        if m_dnbr:  summary_items.append(f"DNBR min={m_dnbr.group(1)} at t={m_dnbr.group(2)} s")

        m_fb    = _re.search(r"Film boiling.*?at t=([\d.]+)\s*s", log)
        if m_fb:    summary_items.append(f"🔥 Film boiling t={m_fb.group(1)} s")

        m_rw    = _re.search(r"Rewetting.*?at t=([\d.]+)\s*s", log)
        if m_rw:    summary_items.append(f"💧 Rewet t={m_rw.group(1)} s")

        # Rod failures from output file
        _run_df = load_results(st.session_state.get("last_case", selected),
                               run_dir=st.session_state.get("run_dir"))
        if _run_df is not None:
            _fail_parts = []
            for _fc, _fl in [
                ("Rod Failures DNB (est.)",    "DNB"),
                ("Rod Failures PCT (est.)",    "PCT"),
                ("Rod Failures Gap (est.)",    "Gap"),
                ("Rod Failures EarlyIV (est.)","EarlyIV"),
            ]:
                if _fc in _run_df.columns:
                    _n = int(_run_df[_fc].max())
                    if _n > 0:
                        _fail_parts.append(f"{_fl}: {_n:,}")
            if _fail_parts:
                summary_items.append("🔴 " + "  ".join(_fail_parts))

        if summary_items:
            st.markdown("**Run summary:** " + "  ·  ".join(summary_items))

        st.markdown("**Console output**")
        st.markdown(f'<div class="console">{log}</div>', unsafe_allow_html=True)


# ── Tab: Results ──────────────────────────────────────────────────────────────
with tab_results:
    run_case = st.session_state.get("last_case", selected)
    _run_dir = st.session_state.get("run_dir")
    df = load_results(run_case, run_dir=_run_dir)

    if df is None:
        st.info("No results yet. Run the simulation first.")
    elif "Time (s)" not in df.columns or len(df) < 2:
        st.warning("Results file exists but contains no simulation data. "
                   "The simulation may have failed  -  check the console output in the Run tab.")
    else:
        # Issue #3: unit system selector (mirrors flare_ua.py)
        unit_sys = st.radio(
            "Units",
            ["Metric  (°C, kPa, kg/s)", "English  (°F, psia, lb/s)"],
            horizontal=True, key="unit_sys_sim",
        )
        _use_english = unit_sys.startswith("English")

        # Scalar summary tiles
        diag_df = load_diagnostics(run_case, run_dir=_run_dir)
        metrics = scalar_summary(
            df,
            use_english=_use_english,
            init_power_mw=st.session_state.get("run_init_power"),
            diag_df=diag_df,
        )
        if metrics:
            html_tiles = '<div class="metric-grid">'
            for item in metrics:
                val, lbl, cls = item[0], item[1], item[2]
                small = len(item) > 3 and item[3]
                v_cls = "val small" if small else "val"
                l_cls = "lbl small" if small else "lbl"
                html_tiles += (f'<div class="metric-tile {cls}">'
                               f'<div class="{v_cls}">{val}</div>'
                               f'<div class="{l_cls}">{lbl}</div>'
                               f'</div>')
            html_tiles += '</div>'
            st.markdown(html_tiles, unsafe_allow_html=True)


        # Plots
        try:
            figs = make_plots(df, run_case, use_english=_use_english)
        except Exception as _me:
            st.error(f"Plot generation error: {_me}")
            import traceback; st.code(traceback.format_exc())
            figs = {}

        all_keys  = [f"fig{i}" for i in range(1, 18)]
        for _fk in ["fig_pzr", "fig18_zr_ecr", "fig19_h2_generation"]:
            if _fk in figs:
                all_keys.append(_fk)
        left_keys  = all_keys[::2]
        right_keys = all_keys[1::2]

        col_l, col_r = st.columns(2)
        for keys, col in [(left_keys, col_l), (right_keys, col_r)]:
            with col:
                for k in keys:
                    if k in figs:
                        try:
                            st.plotly_chart(figs[k], width="stretch",
                                            config={"displayModeBar": False})
                        except Exception as _pe:
                            st.warning(f"Plot {k} error: {_pe}")

        # ── AI Event Narrative ────────────────────────────────────────────────
        st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)
        with st.expander("🤖  AI Event Narrative", expanded=True):
            # Re-read the key each render so adding/editing flare_config.txt
            # during a FLARE session is recognized after a normal Streamlit rerun.
            _current_api_key = load_api_key()
            if not _current_api_key:
                st.warning(
                    "No API key found. Add `ANTHROPIC_API_KEY = sk-ant-...` "
                    "to `runtime/flare_config.txt` (or `Runtime/flare_config.txt`) in the FLARE folder."
                )
            else:
                # Build a structured data summary to ground the narrative
                def _build_narrative_prompt(df, case_name):
                    lines = [f"Case: {case_name}"]

                    t  = df["Time (s)"]
                    dt = t.diff().fillna(0)
                    t0, tf = float(t.iloc[0]), float(t.iloc[-1])
                    lines.append(f"Simulation span: {t0:.1f} – {tf:.1f} s")

                    if "RCS Pressure (kPa)" in df.columns:
                        p = df["RCS Pressure (kPa)"]
                        lines.append(f"RCS pressure: initial {p.iloc[0]:.0f} kPa, "
                                     f"min {p.min():.0f} kPa at t={t.iloc[p.idxmin()]:.1f} s, "
                                     f"final {p.iloc[-1]:.0f} kPa")

                    if "RCS Temperature (K)" in df.columns:
                        T = df["RCS Temperature (K)"] - 273.15
                        lines.append(f"RCS temperature: initial {T.iloc[0]:.1f} °C, "
                                     f"peak {T.max():.1f} °C at t={t.iloc[T.idxmax()]:.1f} s")

                    if "RK Total Power (MW)" in df.columns:
                        pw = df["RK Total Power (MW)"]
                        lines.append(f"Core power: initial {pw.iloc[0]:.1f} MW, "
                                     f"peak {pw.max():.1f} MW, "
                                     f"final {pw.iloc[-1]:.3f} MW")

                    if "Vessel Level (m)" in df.columns:
                        ll = df["Vessel Level (m)"]
                        lines.append(f"Vessel level: initial {ll.iloc[0]:.2f} m, "
                                     f"min {ll.min():.2f} m at t={t.iloc[ll.idxmin()]:.1f} s, "
                                     f"final {ll.iloc[-1]:.2f} m")

                    if "Pump Mass Flow (kg/s)" in df.columns:
                        mf  = df["Pump Mass Flow (kg/s)"]
                        mf0 = float(mf.iloc[0])
                        lines.append(f"Pump flow: initial {mf0:.0f} kg/s, "
                                     f"final {mf.iloc[-1]:.1f} kg/s")
                        if mf0 > 0:
                            trip_idx = (mf < 0.01 * mf0).idxmax()
                            if mf.iloc[trip_idx] < 0.01 * mf0:
                                lines.append(f"RCP flow reaches zero at t={t.iloc[trip_idx]:.1f} s")

                    # Accumulator — enriched
                    if "Accumulator Flow (kg/s)" in df.columns:
                        acc = df["Accumulator Flow (kg/s)"]
                        if acc.max() > 0:
                            active = acc > 0.1
                            t_start = float(t[active].iloc[0])  if active.any() else None
                            t_end   = float(t[active].iloc[-1]) if active.any() else None
                            duration = (t_end - t_start) if (t_start and t_end) else 0
                            total_kg = float((acc * dt).sum())
                            depleted = False
                            if "Accumulator Level (m)" in df.columns:
                                al = df["Accumulator Level (m)"]
                                depleted = bool(al[active.shift(-1, fill_value=False) == False].min() < 0.05) if active.any() else False
                            lines.append(
                                f"Accumulator injection: start t={t_start:.1f} s, "
                                f"end t={t_end:.1f} s, duration {duration:.0f} s, "
                                f"peak {acc.max():.0f} kg/s, "
                                f"total injected ≈{total_kg:.0f} kg"
                                + (", accumulator FULLY DEPLETED" if depleted else "")
                            )

                    # LPSI — enriched
                    if "LPSI Flow (kg/s)" in df.columns:
                        lpsi = df["LPSI Flow (kg/s)"]
                        if lpsi.max() > 0:
                            active = lpsi > 0.1
                            t_start = float(t[active].iloc[0]) if active.any() else None
                            t_end   = float(t[active].iloc[-1]) if active.any() else None
                            total_kg = float((lpsi * dt).sum())
                            lines.append(
                                f"LPSI injection: start t={t_start:.1f} s, "
                                f"end t={t_end:.1f} s, "
                                f"peak {lpsi.max():.0f} kg/s, "
                                f"total injected ≈{total_kg:.0f} kg"
                            )

                    # HPSI — enriched
                    if "HPSI Flow (kg/s)" in df.columns:
                        hpsi = df["HPSI Flow (kg/s)"]
                        if hpsi.max() > 0:
                            active = hpsi > 0.1
                            t_start = float(t[active].iloc[0]) if active.any() else None
                            t_end   = float(t[active].iloc[-1]) if active.any() else None
                            total_kg = float((hpsi * dt).sum())
                            lines.append(
                                f"HPSI injection: start t={t_start:.1f} s, "
                                f"end t={t_end:.1f} s, "
                                f"peak {hpsi.max():.0f} kg/s, "
                                f"total injected ≈{total_kg:.0f} kg"
                            )

                    # CVCS
                    if "CVCS Makeup (kg/s)" in df.columns:
                        mu = df["CVCS Makeup (kg/s)"]
                        if mu.max() > 0:
                            active = mu > 0.1
                            t_start = float(t[active].iloc[0]) if active.any() else None
                            lines.append(f"CVCS makeup/letdown: activates t={t_start:.1f} s, "
                                         f"flow {mu.max():.1f} kg/s")
                            if "RCS Boron (ppm)" in df.columns:
                                bn = df["RCS Boron (ppm)"]
                                lines.append(f"RCS boron: initial {bn.iloc[0]:.0f} ppm, "
                                             f"final {bn.iloc[-1]:.0f} ppm "
                                             f"(Δ{bn.iloc[-1]-bn.iloc[0]:+.0f} ppm)")

                    if "DNBR" in df.columns:
                        dn = df["DNBR"].replace(0, np.nan)
                        mdnbr = dn.min()
                        if not np.isnan(mdnbr):
                            lines.append(f"MDNBR: {mdnbr:.3f} at t={t.iloc[dn.idxmin()]:.1f} s"
                                         + (" — DNB OCCURRED" if mdnbr < 1.0 else ""))

                    if "Clad Surface Temp (K)" in df.columns:
                        tc = df["Clad Surface Temp (K)"] - 273.15
                        lines.append(f"Avg clad temp: initial {tc.iloc[0]:.1f} °C, "
                                     f"peak {tc.max():.1f} °C at t={t.iloc[tc.idxmax()]:.1f} s")

                    if "Hot Pin Clad Temp (K)" in df.columns:
                        thc = df["Hot Pin Clad Temp (K)"] - 273.15
                        lines.append(f"Hot pin clad temp: peak {thc.max():.1f} °C "
                                     f"at t={t.iloc[thc.idxmax()]:.1f} s")

                    if "Hot Pin Fuel Temp (K)" in df.columns:
                        thf = df["Hot Pin Fuel Temp (K)"] - 273.15
                        lines.append(f"Hot pin fuel temp: peak {thf.max():.1f} °C "
                                     f"at t={t.iloc[thf.idxmax()]:.1f} s")

                    if "Clad HTC (W/m2-K)" in df.columns:
                        htc = df["Clad HTC (W/m2-K)"]
                        lines.append(f"Core HTC: initial {htc.iloc[0]:.0f} W/m²K, "
                                     f"min {htc.min():.0f} W/m²K at t={t.iloc[htc.idxmin()]:.1f} s")

                    if "SG Heat Removal (MW)" in df.columns:
                        sg = df["SG Heat Removal (MW)"]
                        if sg.max() > 0:
                            lines.append(f"SG heat removal: peak {sg.max():.1f} MW "
                                         f"at t={t.iloc[sg.idxmax()]:.1f} s")

                    if "PORV Mass Flow (kg/s)" in df.columns:
                        pv = df["PORV Mass Flow (kg/s)"]
                        if pv.max() > 0:
                            active = pv > 0.01
                            t_start = float(t[active].iloc[0]) if active.any() else None
                            t_end   = float(t[active].iloc[-1]) if active.any() else None
                            lines.append(f"PORV: opens t={t_start:.1f} s, "
                                         f"last open t={t_end:.1f} s, "
                                         f"peak {pv.max():.1f} kg/s")

                    if "Pressurizer Level (m)" in df.columns:
                        pzl = df["Pressurizer Level (m)"].replace(0, np.nan)
                        if pzl.notna().any():
                            lines.append(f"Pressurizer level: initial {pzl.iloc[0]:.2f} m, "
                                         f"min {pzl.min():.2f} m at t={t.iloc[pzl.idxmin()]:.1f} s, "
                                         f"final {pzl.iloc[-1]:.2f} m")

                    for col, label in [
                        ("Rod Failures DNB (est.)",     "Rod failures (DNB)"),
                        ("Rod Failures Gap (est.)",     "Rod failures (gap release)"),
                        ("Rod Failures EarlyIV (est.)", "Rod failures (early in-vessel)"),
                    ]:
                        if col in df.columns and df[col].max() > 0:
                            lines.append(f"{label}: {int(df[col].max())} rods")

                    if "Equilibrium Quality (-)" in df.columns:
                        xeq = df["Equilibrium Quality (-)"]
                        lines.append(f"Equilibrium quality: min {xeq.min():.3f}, "
                                     f"max {xeq.max():.3f}")

                    # Narrative scope and acceptance-criteria guardrails
                    try:
                        _pwr0 = float(df["RK Total Power (MW)"].iloc[0]) if "RK Total Power (MW)" in df.columns else 0.0
                        _sg0 = float(df["SG Heat Removal (MW)"].iloc[0]) if "SG Heat Removal (MW)" in df.columns else 0.0
                        if _pwr0 > 0 and _sg0 < 0.5 * _pwr0:
                            lines.append("Initial-condition check: initial SG heat removal is less than 50% of initial reactor power; do not describe this as normal full-power steady-state operation unless other data support that conclusion.")
                    except Exception:
                        pass

                    try:
                        _pct_limit_K = 1477.0
                        _pct = None
                        if "Hot Pin Clad Temp (K)" in df.columns:
                            _pct = float(pd.to_numeric(df["Hot Pin Clad Temp (K)"], errors="coerce").max())
                            if _pct >= _pct_limit_K:
                                lines.append(f"10 CFR 50.46 PCT check: peak hot-pin clad temperature {_pct:.1f} K exceeds the 1477 K limit; state this as a limit exceedance.")
                            else:
                                lines.append(f"10 CFR 50.46 PCT check: peak hot-pin clad temperature {_pct:.1f} K is below the 1477 K limit.")
                    except Exception:
                        pass

                    try:
                        _ecr = float(pd.to_numeric(df.get("Zr Oxidation Mean Oxidizing Rod ECR (%)", pd.Series([0.0])), errors="coerce").fillna(0.0).max())
                        _h2 = float(pd.to_numeric(df.get("H2 Generated (kg)", pd.Series([0.0])), errors="coerce").fillna(0.0).max())
                        _rod_cols = ["Rod Failures DNB (est.)", "Rod Failures Gap (est.)", "Rod Failures EarlyIV (est.)"]
                        _rods = 0.0
                        for _c in _rod_cols:
                            if _c in df.columns:
                                _rods = max(_rods, float(pd.to_numeric(df[_c], errors="coerce").fillna(0.0).max()))
                        if _ecr <= 0.0 and _h2 <= 0.0 and _rods <= 0.0:
                            lines.append("Fuel-integrity scope: ECR, hydrogen generation, and rod failures are zero; collapse fuel-integrity discussion to one concise sentence.")
                        if (case_name or "").lower().endswith("null") or (_ecr <= 0.0 and _h2 <= 0.0 and _rods <= 0.0 and _pct is not None and _pct < 900.0):
                            lines.append("Narrative scope: treat this as a verification/null or benign case; use a short three-section structure rather than a full accident narrative.")
                    except Exception:
                        pass

                    try:
                        if "DNBR" in df.columns:
                            _dn_final = pd.to_numeric(df["DNBR"], errors="coerce").iloc[-1]
                            if pd.isna(_dn_final):
                                lines.append("DNBR interpretation: final DNBR is NaN because the DNB correlation is not applicable under the final conditions; this is expected behavior, not a numerical convergence issue.")
                    except Exception:
                        pass

                    _st_lines = _source_term_lines_for_narrative(case_name, run_dir=_run_dir)
                    if _st_lines:
                        lines.append("")
                        lines.append("SOURCE TERM AND RADIOLOGICAL DATA FOR NARRATIVE:")
                        lines.extend(_st_lines)

                    return "\n".join(lines)

                def _build_sequence_of_events(df, filter_keywords=None):
                    """Detect key timestamped events from the data for the SOE table."""
                    t   = df["Time (s)"]
                    dt  = t.diff().fillna(0)
                    soe = []

                    def first_time(mask):
                        idx = mask.idxmax() if mask.any() else None
                        return float(t.iloc[idx]) if idx is not None and mask.any() else None

                    def last_time(mask):
                        idxs = t[mask]
                        return float(idxs.iloc[-1]) if len(idxs) > 0 else None

                    # Simulation start
                    soe.append((float(t.iloc[0]), "Simulation start — reactor at full power"))

                    # Reactor power events
                    if "RK Total Power (MW)" in df.columns:
                        pw  = df["RK Total Power (MW)"]
                        pw0 = float(pw.iloc[0])
                        op  = first_time(pw > 1.05 * pw0)
                        scr = first_time((t > float(t.iloc[0]) + 0.5) & (pw < 0.5 * pw0))
                        if op is not None and (scr is None or op < scr):
                            soe.append((op, f"High-flux trip setpoint reached ({pw.max():.0f} MW peak)"))
                        if scr is not None:
                            soe.append((scr, f"Reactor scram — power below 50% rated"))

                    # Pump coastdown
                    if "Pump Mass Flow (kg/s)" in df.columns:
                        mf  = df["Pump Mass Flow (kg/s)"]
                        mf0 = float(mf.iloc[0])
                        if mf0 > 0:
                            t50 = first_time((t > 0.5) & (mf < 0.5 * mf0))
                            t00 = first_time((t > 0.5) & (mf < 0.01 * mf0))
                            if t50: soe.append((t50, "RCP coastdown — pump flow below 50% rated"))
                            if t00: soe.append((t00, "RCP flow effectively zero"))

                    # PORV
                    if "PORV Mass Flow (kg/s)" in df.columns:
                        pv = df["PORV Mass Flow (kg/s)"]
                        if pv.max() > 0.01:
                            tp = first_time(pv > 0.01)
                            tc = last_time(pv > 0.01)
                            if tp: soe.append((tp, f"PORV opens ({pv.max():.1f} kg/s peak)"))
                            if tc and tc != tp:
                                soe.append((tc, "PORV last open — RCS pressure stabilising"))

                    # RCS pressure min
                    if "RCS Pressure (kPa)" in df.columns:
                        p = df["RCS Pressure (kPa)"]
                        soe.append((float(t.iloc[p.idxmin()]),
                                    f"Minimum RCS pressure — {p.min():.0f} kPa"))

                    # Accumulator — start, depletion, end
                    if "Accumulator Flow (kg/s)" in df.columns:
                        acc    = df["Accumulator Flow (kg/s)"]
                        active = acc > 0.1
                        if active.any():
                            ta_start = first_time(active)
                            ta_end   = last_time(active)
                            total_kg = float((acc * dt).sum())
                            soe.append((ta_start,
                                        f"Accumulator injection begins ({acc.max():.0f} kg/s peak)"))
                            # Depletion: level reaches near-zero
                            depleted = False
                            if "Accumulator Level (m)" in df.columns:
                                al = df["Accumulator Level (m)"]
                                depl_mask = active & (al < 0.05)
                                if depl_mask.any():
                                    td = first_time(depl_mask)
                                    soe.append((td,
                                        f"Accumulator empty — injection ends "
                                        f"(≈{total_kg:.0f} kg total injected)"))
                                    depleted = True
                            if not depleted and ta_end and ta_end != ta_start:
                                soe.append((ta_end,
                                    f"Accumulator injection ends "
                                    f"(≈{total_kg:.0f} kg total injected)"))

                    # LPSI
                    if "LPSI Flow (kg/s)" in df.columns:
                        lpsi   = df["LPSI Flow (kg/s)"]
                        active = lpsi > 0.1
                        if active.any():
                            total_kg = float((lpsi * dt).sum())
                            soe.append((first_time(active),
                                        f"LPSI injection begins ({lpsi.max():.0f} kg/s peak)"))
                            tl_end = last_time(active)
                            if tl_end and tl_end != first_time(active):
                                soe.append((tl_end,
                                            f"LPSI injection ends (≈{total_kg:.0f} kg total)"))

                    # HPSI
                    if "HPSI Flow (kg/s)" in df.columns:
                        hpsi   = df["HPSI Flow (kg/s)"]
                        active = hpsi > 0.1
                        if active.any():
                            total_kg = float((hpsi * dt).sum())
                            soe.append((first_time(active),
                                        f"HPSI injection begins ({hpsi.max():.0f} kg/s peak)"))

                    # CVCS
                    if "CVCS Makeup (kg/s)" in df.columns:
                        mu = df["CVCS Makeup (kg/s)"]
                        if mu.max() > 0.1:
                            soe.append((first_time(mu > 0.1),
                                        f"CVCS makeup/letdown activates ({mu.max():.1f} kg/s)"))

                    # Boron dilution trip
                    if "RCS Boron (ppm)" in df.columns:
                        bn  = df["RCS Boron (ppm)"]
                        bn0 = float(bn.iloc[0])
                        if bn0 > 0 and bn.min() < 0.95 * bn0:
                            td = first_time(bn < 0.95 * bn0)
                            soe.append((td,
                                        f"RCS boron below 95% initial "
                                        f"({bn0:.0f}→{bn.min():.0f} ppm)"))

                    # DNB
                    if "DNBR" in df.columns:
                        dn = df["DNBR"].replace(0, np.nan)
                        if dn.min() < 1.0:
                            soe.append((first_time(dn < 1.0),
                                        f"DNB — DNBR drops below 1.0 (min {dn.min():.3f})"))

                    # Vessel level min and recovery
                    if "Vessel Level (m)" in df.columns:
                        ll   = df["Vessel Level (m)"]
                        ll0  = float(ll.iloc[0])
                        tidx = ll.idxmin()
                        soe.append((float(t.iloc[tidx]),
                                    f"Minimum vessel level — {ll.min():.2f} m"))
                        # Recovery: level recovers to 50% of initial after minimum
                        recov = (t > float(t.iloc[tidx])) & (ll > 0.5 * ll0)
                        tr = first_time(recov)
                        if tr:
                            soe.append((tr, f"Vessel level recovers to 50% initial ({0.5*ll0:.2f} m)"))

                    # Peak fuel temperature
                    if "Hot Pin Fuel Temp (K)" in df.columns:
                        thf = df["Hot Pin Fuel Temp (K)"] - 273.15
                        soe.append((float(t.iloc[thf.idxmax()]),
                                    f"Peak hot pin fuel temperature — {thf.max():.0f} °C"))

                    # Peak clad temperature
                    if "Hot Pin Clad Temp (K)" in df.columns:
                        thc = df["Hot Pin Clad Temp (K)"] - 273.15
                        soe.append((float(t.iloc[thc.idxmax()]),
                                    f"Peak hot pin clad temperature — {thc.max():.0f} °C"))
                    elif "Clad Surface Temp (K)" in df.columns:
                        tc3 = df["Clad Surface Temp (K)"] - 273.15
                        soe.append((float(t.iloc[tc3.idxmax()]),
                                    f"Peak average clad temperature — {tc3.max():.0f} °C"))

                    # Pressurizer level min
                    if "Pressurizer Level (m)" in df.columns:
                        pzl = df["Pressurizer Level (m)"].replace(0, np.nan)
                        if pzl.notna().any() and pzl.min() < float(pzl.iloc[0]) * 0.9:
                            soe.append((float(t.iloc[pzl.idxmin()]),
                                        f"Minimum pressurizer level — {pzl.min():.2f} m"))

                    # End of simulation
                    soe.append((float(t.iloc[-1]), "End of simulation"))

                    # Sort and deduplicate (remove identical times keeping first)
                    soe.sort(key=lambda x: x[0])
                    seen = set()
                    soe_dedup = []
                    for ts, desc in soe:
                        key = round(ts, 1)
                        if key not in seen:
                            seen.add(key)
                            soe_dedup.append((ts, desc))
                    soe = soe_dedup

                    # Filter for Brief mode
                    if filter_keywords:
                        soe = [(ts, desc) for ts, desc in soe
                               if any(kw.lower() in desc.lower()
                                      for kw in filter_keywords)]

                    # Build markdown table
                    table  = "\n\n---\n\n**Sequence of Events**\n\n"
                    table += "| Time (s) | Event |\n"
                    table += "|---:|:------|\n"
                    for ts, desc in soe:
                        table += f"| {ts:.1f} | {desc} |\n"
                    return table
                    table += "|---:|:------|\n"
                    for ts, desc in soe:
                        table += f"| {ts:.1f} | {desc} |\n"
                    return table

                _narrative_key = f"narrative_{run_case}"
                _tables_key    = f"tables_{run_case}"
                _tables_only_key = f"tables_only_{run_case}"
                if _narrative_key not in st.session_state:
                    st.session_state[_narrative_key] = None
                if _tables_key not in st.session_state:
                    st.session_state[_tables_key] = None
                if _tables_only_key not in st.session_state:
                    st.session_state[_tables_only_key] = False

                _detail_val = st.slider(
                    "Detail level", min_value=0.0, max_value=1.0, value=0.5,
                    step=0.05,
                    key="narrative_detail",
                    help="0 = Brief summary  |  0.5 = Standard  |  1.0 = Comprehensive analysis",
                )
                # Map continuous 0–1 to label and parameters
                if _detail_val < 0.25:
                    _detail_label = "Brief"
                elif _detail_val < 0.75:
                    _detail_label = "Standard"
                else:
                    _detail_label = "Detailed"
                st.caption(f"**{_detail_label}**")

                # Interpolate max_tokens: 300 at 0 → 2048 at 1
                _max_tokens = int(300 + _detail_val * (2048 - 300))

                # Interpolate paragraph count
                _para_min = max(1, round(1 + _detail_val * 7))   # 1 at 0 → 8 at 1
                _para_max = _para_min + 1
                _paragraphs = (f"{_para_min} paragraph" if _para_min == 1
                               else f"{_para_min}–{_para_max} paragraphs")

                # Instruction scales continuously — more topics added as detail increases
                _topics = [
                    (0.00, "the initiating event and immediate plant response"),
                    (0.20, "the thermal-hydraulic progression"),
                    (0.40, "safety system performance (ECCS, PORV, RPS)"),
                    (0.55, "fuel and cladding thermal response with peak values"),
                    (0.65, "reactivity feedback and kinetics behaviour"),
                    (0.75, "acceptance criteria evaluation and regulatory significance"),
                    (0.85, "specific parameter values and times for all key events"),
                    (0.92, "uncertainty considerations and sensitivity to key assumptions"),
                ]
                _active = [t for thresh, t in _topics if _detail_val >= thresh]
                if len(_active) == 1:
                    _instruction = f"Focus on {_active[0]}. Omit minor details."
                else:
                    _instruction = ("Cover the following in order: "
                                    + "; ".join(f"({i+1}) {t}" for i, t in enumerate(_active)) + ".")

                # SOE filter: brief mode only includes key events
                _soe_filter = ({"scram", "trip", "inject", "DNB", "peak", "End", "start"}
                                if _detail_val < 0.25 else None)

                def _build_ic_final(df):
                        """Build a list of (parameter, initial, final, units) rows."""
                        t   = df["Time (s)"]
                        rows = []

                        def ic_fin(col, label, units, scale=1.0, fmt=".1f", offset=0.0):
                            if col in df.columns:
                                s = df[col].dropna()
                                if len(s) < 2:
                                    return
                                ic  = (float(s.iloc[0])  + offset) * scale
                                fin = (float(s.iloc[-1]) + offset) * scale
                                rows.append((label,
                                             f"{ic:{fmt}}",
                                             f"{fin:{fmt}}",
                                             units))

                        ic_fin("RCS Pressure (kPa)",       "RCS Pressure",          "kPa",  fmt=".0f")
                        ic_fin("RCS Temperature (K)",      "RCS Temperature",       "°C",   offset=-273.15, fmt=".1f")
                        ic_fin("RK Total Power (MW)",      "Core Power",            "MW",   fmt=".1f")
                        ic_fin("Core Power (MW)",          "Heat to Coolant",       "MW",   fmt=".1f")
                        ic_fin("Total Mass Scaled",        "RCS Inventory (scaled)","—",    fmt=".4f")
                        ic_fin("Vessel Level (m)",         "Vessel Level",          "m",    fmt=".2f")
                        ic_fin("Pump Mass Flow (kg/s)",    "RCP Mass Flow",         "kg/s", fmt=".0f")
                        ic_fin("Pump Speed (rpm)",         "RCP Speed",             "rpm",  fmt=".0f")
                        ic_fin("Clad Surface Temp (K)",    "Avg Clad Temp",         "°C",   offset=-273.15, fmt=".1f")
                        ic_fin("Hot Pin Clad Temp (K)",    "Hot Pin Clad Temp",     "°C",   offset=-273.15, fmt=".1f")
                        ic_fin("Hot Pin Fuel Temp (K)",    "Hot Pin Fuel Temp",     "°C",   offset=-273.15, fmt=".1f")
                        ic_fin("DNBR",                     "DNBR",                  "—",    fmt=".3f")
                        ic_fin("Equilibrium Quality (-)",  "Equilibrium Quality",   "—",    fmt=".4f")
                        ic_fin("Void Fraction (-)",        "Void Fraction",         "—",    fmt=".4f")
                        ic_fin("SG Heat Removal (MW)",     "SG Heat Removal",       "MW",   fmt=".1f")
                        ic_fin("Pressurizer Level (m)",    "Pressurizer Level",     "m",    fmt=".2f")
                        ic_fin("RCS Boron (ppm)",          "RCS Boron",             "ppm",  fmt=".0f")
                        ic_fin("Accumulator Level (m)",    "Accumulator Level",     "m",    fmt=".2f")
                        ic_fin("Reactivity net (pcm)",     "Net Reactivity",        "pcm",  fmt=".1f")
                        return rows

                _btn_col1, _btn_col2 = st.columns([1, 1])
                with _btn_col1:
                    _gen_narrative = st.button("Generate narrative",
                                               key="gen_narrative", type="primary",
                                               help="Requires Anthropic API key")
                with _btn_col2:
                    _gen_tables = st.button("Build tables only",
                                            key="gen_tables",
                                            help="No API key needed — builds IC, SOE and Word download locally")

                # Tables-only path (no API call)
                if _gen_tables:
                    _soe_tbl   = _build_sequence_of_events(df, filter_keywords=None)
                    _ic_r      = _build_ic_final(df)
                    st.session_state[_tables_key] = {"soe": _soe_tbl, "ic": _ic_r, "tables_only": True}

                # Full narrative path (API call)
                if _gen_narrative:
                    with st.spinner("Analysing results…"):
                        _data_summary = _build_narrative_prompt(df, run_case)
                        _soe_table    = _build_sequence_of_events(
                            df, filter_keywords=_soe_filter
                        )
                        _prompt = (
                            "You are a PWR safety analysis engineer reviewing the output "
                            "of a FLARE transient simulation. Write a clear, concise "
                            f"technical narrative ({_paragraphs}) describing "
                            "the event sequence, the plant response, and the safety significance. "
                            "Use standard nuclear engineering terminology. "
                            f"{_instruction} "
                            + NARRATIVE_QUALITY_GUIDANCE + " "
                            "Use explicit section headings. Keep fuel integrity/core-damage topics "
                            "separate from source-term and radiological topics. Discuss ECR, hydrogen generation, "
                            "DNB, cladding temperature, and rod failures only under a heading such as "
                            "'Fuel Integrity and Core Damage'. Discuss radionuclide release fractions, source-term model, "
                            "NOTBADTRAD dose, iodine-spike dose, and radiological consequences only under "
                            "'Source Term and Radiological Consequences'. Do not merge those topics. "
                            "If source-term data are supplied, include the important numerical group-release and dose results "
                            "from that data in the source-term section. "
                            "Base your narrative strictly on the data provided — "
                            "do not invent values not present in the summary. "
                            "End your response after the final paragraph — do not "
                            "add a sequence of events table (one is appended separately).\n\n"
                            "Simulation data summary:\n" + _data_summary
                        )
                        try:
                            _resp = requests.post(
                                "https://api.anthropic.com/v1/messages",
                                headers={
                                    "x-api-key":         _current_api_key,
                                    "anthropic-version": "2023-06-01",
                                    "content-type":      "application/json",
                                },
                                json={
                                    "model":      _ANTHROPIC_MODEL,
                                    "max_tokens": _max_tokens,
                                    "messages":   [{"role": "user", "content": _prompt}],
                                },
                                timeout=60,
                            )
                            _resp.raise_for_status()
                            st.session_state[_narrative_key] = (
                                _resp.json()["content"][0]["text"] + _soe_table
                            )
                            # Also store SOE so download works; IC/final-state table is
                            # displayed only when the user selects Build tables only.
                            st.session_state[_tables_key] = {
                                "soe": _soe_table,
                                "ic":  _build_ic_final(df),
                                "tables_only": False,
                            }
                        except Exception as _ne:
                            st.session_state[_narrative_key] = f"⚠️  API error: {_ne}"

                # ── Display: narrative (if generated) then tables ─────────────
                if st.session_state.get(_narrative_key):
                    st.markdown(st.session_state[_narrative_key])

                # Tables come from either path. The Initial Conditions & Final State
                # table is intentionally displayed only for the explicit
                # "Build tables only" path; generated narratives keep the run view concise.
                _stored_tables = st.session_state.get(_tables_key)
                _tables_only_active = bool(_stored_tables and _stored_tables.get("tables_only"))
                _ic_rows = (_stored_tables.get("ic", [])
                            if (_stored_tables and _tables_only_active) else [])

                # SOE — render explicitly when no narrative (tables-only path)
                if _stored_tables and not st.session_state.get(_narrative_key):
                    _soe_md = _stored_tables.get("soe", "")
                    if _soe_md:
                        st.markdown(_soe_md)

                if _tables_only_active:
                    if not _ic_rows:
                        _ic_rows = _build_ic_final(df)
                    if _ic_rows:
                        st.markdown("---\n**Initial Conditions & Final State**")
                        import pandas as _pd2
                        _ic_df = _pd2.DataFrame(
                            _ic_rows, columns=["Parameter", "Initial", "Final", "Units"]
                        )
                        st.dataframe(_ic_df, hide_index=True, width="stretch")

                # ── Download as Word document (available from either button) ─
                if _stored_tables or st.session_state.get(_narrative_key):
                    def _narrative_to_docx(text, case_name, ic_rows=None, report_tables=None):
                        """Convert narrative markdown + SOE/source-term/dose tables to a .docx bytes buffer."""
                        try:
                            from docx import Document as DocxDocument
                            from docx.shared import Pt, RGBColor, Inches
                        except ImportError:
                            return None   # python-docx not installed

                        import re, io
                        from datetime import date

                        doc = DocxDocument()

                        # Page margins
                        for section in doc.sections:
                            section.top_margin    = Inches(1)
                            section.bottom_margin = Inches(1)
                            section.left_margin   = Inches(1.25)
                            section.right_margin  = Inches(1.25)

                        # Title
                        title = doc.add_heading(f"Event Narrative — {case_name}", level=1)
                        title.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                        # Date line
                        from datetime import date
                        dp = doc.add_paragraph(f"Generated: {date.today().isoformat()}")
                        dp.runs[0].font.size = Pt(9)
                        dp.runs[0].font.color.rgb = RGBColor(0x80, 0x80, 0x80)
                        doc.add_paragraph()   # spacer

                        # Split narrative from SOE table
                        parts = text.split("\n\n---\n\n")
                        narrative_text = parts[0].strip()
                        soe_text = parts[1].strip() if len(parts) > 1 else None

                        # Write narrative paragraphs
                        for para in narrative_text.split("\n\n"):
                            para = para.strip()
                            if not para:
                                continue
                            # Bold inline (**text**)
                            p = doc.add_paragraph()
                            p.paragraph_format.space_after = Pt(6)
                            # Simple inline bold parse
                            segments = re.split(r'\*\*(.*?)\*\*', para)
                            for i, seg in enumerate(segments):
                                run = p.add_run(seg)
                                run.bold = (i % 2 == 1)
                                run.font.size = Pt(11)

                        # SOE table
                        if soe_text:
                            doc.add_paragraph()
                            h = doc.add_heading("Sequence of Events", level=2)
                            h.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)

                            # Parse markdown table rows
                            rows = [r for r in soe_text.split("\n")
                                    if r.startswith("|") and "---" not in r]
                            if rows:
                                # Extract cells
                                parsed = [[c.strip() for c in r.strip("|").split("|")]
                                          for r in rows]
                                tbl = doc.add_table(rows=len(parsed), cols=2)
                                tbl.style = "Table Grid"
                                for ri, row_data in enumerate(parsed):
                                    cells = tbl.rows[ri].cells
                                    for ci, val in enumerate(row_data[:2]):
                                        cells[ci].text = val
                                        run = cells[ci].paragraphs[0].runs
                                        if run:
                                            run[0].font.size = Pt(10)
                                            if ri == 0:   # header row
                                                run[0].bold = True

                        # IC / Final State table
                        if ic_rows:
                            doc.add_paragraph()
                            h2 = doc.add_heading("Initial Conditions & Final State", level=2)
                            h2.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                            border = {"style": "single", "size": 1, "color": "CCCCCC"}
                            from docx.oxml.ns import qn
                            from docx.oxml import OxmlElement
                            tbl2 = doc.add_table(rows=1 + len(ic_rows), cols=4)
                            tbl2.style = "Table Grid"
                            # Header row
                            hdr_cells = tbl2.rows[0].cells
                            for ci, hdr in enumerate(["Parameter", "Initial", "Final", "Units"]):
                                hdr_cells[ci].text = hdr
                                run = hdr_cells[ci].paragraphs[0].runs
                                if run:
                                    run[0].bold = True
                                    run[0].font.size = Pt(10)
                            # Data rows
                            for ri, (param, ic, fin, units) in enumerate(ic_rows):
                                row_cells = tbl2.rows[ri + 1].cells
                                for ci, val in enumerate([param, ic, fin, units]):
                                    row_cells[ci].text = val
                                    run = row_cells[ci].paragraphs[0].runs
                                    if run:
                                        run[0].font.size = Pt(10)


                        # Source-term and dose tables exactly as displayed in Results panel
                        if report_tables:
                            current_section = None
                            for item in report_tables:
                                section = item.get("section", "")
                                title_txt = item.get("title", "")
                                df_tbl = item.get("df")
                                if df_tbl is None or getattr(df_tbl, "empty", True):
                                    continue
                                if section != current_section:
                                    doc.add_paragraph()
                                    h3 = doc.add_heading(section, level=2)
                                    h3.runs[0].font.color.rgb = RGBColor(0x2E, 0x75, 0xB6)
                                    current_section = section
                                if title_txt:
                                    p_title = doc.add_paragraph()
                                    r_title = p_title.add_run(title_txt)
                                    r_title.bold = True
                                    r_title.font.size = Pt(10)
                                cols = [str(c) for c in df_tbl.columns]
                                tblx = doc.add_table(rows=1 + len(df_tbl), cols=len(cols))
                                tblx.style = "Table Grid"
                                for ci, hdr in enumerate(cols):
                                    cell = tblx.rows[0].cells[ci]
                                    cell.text = hdr
                                    runs = cell.paragraphs[0].runs
                                    if runs:
                                        runs[0].bold = True
                                        runs[0].font.size = Pt(8)
                                for ri, (_, row) in enumerate(df_tbl.iterrows(), start=1):
                                    for ci, col in enumerate(cols):
                                        cell = tblx.rows[ri].cells[ci]
                                        cell.text = str(row.get(col, ""))
                                        runs = cell.paragraphs[0].runs
                                        if runs:
                                            runs[0].font.size = Pt(8)
                                doc.add_paragraph()

                        buf = io.BytesIO()
                        doc.save(buf)
                        buf.seek(0)
                        return buf.read()

                    _docx_text = (st.session_state.get(_narrative_key) or
                                  (_stored_tables.get("soe", "") if _stored_tables else ""))
                    _docx_report_tables = build_source_term_dose_report_tables(run_case, run_dir=_run_dir)
                    _docx_bytes = _narrative_to_docx(
                        _docx_text, run_case,
                        ic_rows=_ic_rows,
                        report_tables=_docx_report_tables,
                    )
                    if _docx_bytes is not None:
                        st.download_button(
                            label="⬇  Download narrative (.docx)",
                            data=_docx_bytes,
                            file_name=f"{run_case}_narrative.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="dl_narrative_docx",
                        )
                    else:
                        st.info("Install **python-docx** to enable Word download: "
                                "`pip install python-docx`")


        # ── NBT Source Term and Dose Results ───────────────────────────────
        # Kept after the main FLARE nuclear thermal-hydraulic results so the
        # radiological screening is clearly downstream of the transient results.
        # Source term
        st_df = load_source_term(run_case, run_dir=_run_dir)
        if st_df is not None and not st_df.empty:
            with st.expander("🔴  Source Term  -  RG 1.183 / NBT", expanded=True):
                try:
                    _st_option, _st_applied = _extract_source_term_basis(
                        st_df, case_name=run_case, run_dir=_run_dir
                    )
                    if _st_option:
                        if _st_option == "licensing-auto":
                            _applied_txt = _st_applied or "not reported in this output file"
                            st.info(
                                f"**Source-term option:** `{_st_option}`  ·  "
                                f"**Applied model:** `{_applied_txt}`"
                            )
                        else:
                            st.info(f"**Source-term option:** `{_st_option}`")

                    # Main source-term rows are those with a numeric inventory.
                    inv_col = "Core inventory (Ci)"
                    rel_col = "Released inventory (Ci)"
                    frac_col = "Total release frac"

                    if inv_col in st_df.columns:
                        main_mask = st_df[inv_col].apply(
                            lambda v: isinstance(v, (int, float, np.integer, np.floating))
                            and not pd.isna(v)
                        )
                        main_df = st_df[main_mask].copy()
                        main_df = _clean_source_term_display_df(main_df)

                        if not main_df.empty:
                            # Ensure TOTAL appears last and format numeric columns cleanly.
                            display_cols = [c for c in [
                                "Group",
                                "NBT key",
                                inv_col,
                                "Gap release %",
                                "Early IV %",
                                "Total release %",
                                frac_col,
                                rel_col,
                            ] if c in main_df.columns]

                            st.markdown("**Group inventories and releases:**")
                            fmt = {}
                            for c in [inv_col, rel_col]:
                                if c in main_df.columns:
                                    fmt[c] = "{:.4e}"
                            for c in ["Gap release %", "Early IV %", "Total release %"]:
                                if c in main_df.columns:
                                    fmt[c] = "{:.3f}"
                            if frac_col in main_df.columns:
                                fmt[frac_col] = "{:.4e}"

                            st.dataframe(
                                main_df[display_cols].reset_index(drop=True).style.format(fmt),
                                width="stretch", hide_index=True,
                            )

                    # Show model metadata separately, if present.
                    meta_labels = {
                        "Model selected",
                        "Fuel source-term option",
                        "Inventory model",
                        "Estimated fissions",
                        "NOTBADTRAD gap release duration (hr)",
                        "NOTBADTRAD early-IV release duration (hr)",
                        "NOTBADTRAD EAB integration time (hr)",
                        "NOTBADTRAD LPZ integration time (hr)",
                    }
                    if "Group" in st_df.columns:
                        meta_df = st_df[st_df["Group"].isin(meta_labels)].copy()
                        # Metadata rows reuse the Total release frac column as a value field;
                        # it intentionally mixes text labels and numbers, so display it as text.
                        meta_df = _clean_numeric_columns_for_streamlit(
                            meta_df, numeric_cols=[], text_cols=list(meta_df.columns)
                        )
                        if not meta_df.empty:
                            meta_cols = [c for c in ["Group", frac_col] if c in meta_df.columns]
                            if meta_cols:
                                st.markdown("**Source-term basis:**")
                                st.dataframe(
                                    meta_df[meta_cols].reset_index(drop=True),
                                    width="stretch", hide_index=True,
                                )

                    # Backward-compatible fallback for older output files.
                    if inv_col not in st_df.columns:
                        display_cols = [c for c in
                                        ["Group", "Gap release %",
                                         "Early IV %", "Total release %"]
                                        if c in st_df.columns]
                        if display_cols:
                            _fallback_df = _clean_source_term_display_df(st_df[display_cols])
                            st.dataframe(_fallback_df, width="stretch", hide_index=True)

                except Exception as _ste:
                    st.warning(f"Could not parse source-term results: {_ste}")
                    st.dataframe(st_df.astype(str), width="stretch", hide_index=True)

        # ── NOTBADTRAD Dose Screening ─────────────────────────────────────
        dose_rows = load_dose(run_case, run_dir=_run_dir)
        if dose_rows:
            with st.expander("☢️  NOTBADTRAD Dose Screening", expanded=True):
                try:
                    # ── Summary table (EAB / LPZ / Control Room) ──────────
                    summary = [r[:5] for r in dose_rows
                               if r[0] in ("EAB", "LPZ", "Control Room")]
                    if summary:
                        sum_df = pd.DataFrame(summary,
                            columns=["Location", "TEDE (rem)", "Limit (rem)",
                                     "Margin (rem)", "Result"])
                        sum_df = _clean_dose_display_df(sum_df)
                        def _colour_result(v):
                            return ("color: green; font-weight: bold"
                                    if v == "PASS" else
                                    "color: red; font-weight: bold")
                        st.dataframe(
                            sum_df.style
                                .map(_colour_result, subset=["Result"])
                                .format({"TEDE (rem)": "{:.4e}",
                                         "Margin (rem)": "{:.4e}"}),
                            width="stretch", hide_index=True,
                        )

                    # ── Group contributions ────────────────────────────────
                    grp_hdr = next((i for i, r in enumerate(dose_rows)
                                    if r[0] == "Group"), None)
                    if grp_hdr is not None:
                        grp_rows = []
                        for r in dose_rows[grp_hdr + 1:]:
                            if not any(r): break
                            grp_rows.append(r)
                        if grp_rows:
                            grp_df = pd.DataFrame(
                                    [r[:6] for r in grp_rows],
                                    columns=["Group","Released (Ci)","EAB (rem)","LPZ (rem)",
                                             "CR (rem)","Release frac"])
                            grp_df = _clean_dose_display_df(grp_df)
                            st.markdown("**Group contributions:**")
                            st.dataframe(
                                grp_df.style.format(
                                    {"Released (Ci)": "{:.4e}",
                                     "EAB (rem)": "{:.4e}", "LPZ (rem)": "{:.4e}",
                                     "CR (rem)":  "{:.4e}", "Release frac": "{:.4e}"}),
                                width="stretch", hide_index=True,
                            )

                    # ── Runtime ────────────────────────────────────────────
                    rt = next((r[1] for r in dose_rows
                               if r[0] == "Runtime (ms)"), None)
                    if rt is not None:
                        st.caption(f"NOTBADTRAD runtime = {rt:.1f} ms")

                    # ── Power-law fit and distance table ───────────────────
                    fit_row = next((r for r in dose_rows
                                    if r[0] == "Power-law fit"), None)
                    if fit_row and fit_row[1]:
                        st.markdown("---")
                        st.markdown("**TEDE vs Distance  —  EAB/LPZ integration intervals:**")
                        st.code(fit_row[1], language=None)
                        fit_err = next((r[1] for r in dose_rows
                                        if r[0] == "Fit error (%)"), None)
                        if fit_err is not None:
                            st.caption(f"Fit error: {fit_err:.4f}%  "
                                       "(exact for Pasquill-Gifford dispersion)")

                        # Distance table rows
                        dist_hdr = next((i for i, r in enumerate(dose_rows)
                                         if r[0] == "Distance (m)"), None)
                        if dist_hdr is not None:
                            dt_rows = []
                            for r in dose_rows[dist_hdr + 1:]:
                                if not any(r): break
                                dt_rows.append(r)
                            if dt_rows:
                                hdr = dose_rows[dist_hdr]
                                if len(hdr) >= 4 and "chi" in str(hdr[1]).lower() and "EAB" in str(hdr[2]) and "LPZ" in str(hdr[3]):
                                    dt_df = pd.DataFrame(
                                        [r[:4] for r in dt_rows],
                                        columns=["Distance (m)",
                                                 "χ/Q (s/m³)",
                                                 "TEDE EAB interval (rem)",
                                                 "TEDE LPZ interval (rem)"])
                                    dt_df = _clean_dose_display_df(dt_df)
                                    dt_df["Distance (m)"] = dt_df["Distance (m)"].astype("Int64")
                                    st.dataframe(
                                        dt_df.style.format(
                                            {"χ/Q (s/m³)": "{:.4e}",
                                             "TEDE EAB interval (rem)": "{:.4e}",
                                             "TEDE LPZ interval (rem)": "{:.4e}"}),
                                        width="stretch", hide_index=True,
                                    )
                                elif len(hdr) >= 3 and "EAB" in str(hdr[1]) and "LPZ" in str(hdr[2]):
                                    dt_df = pd.DataFrame(
                                        [r[:3] for r in dt_rows],
                                        columns=["Distance (m)",
                                                 "TEDE EAB interval (rem)",
                                                 "TEDE LPZ interval (rem)"])
                                    dt_df = _clean_dose_display_df(dt_df)
                                    dt_df["Distance (m)"] = dt_df["Distance (m)"].astype("Int64")
                                    st.dataframe(
                                        dt_df.style.format(
                                            {"TEDE EAB interval (rem)": "{:.4e}",
                                             "TEDE LPZ interval (rem)": "{:.4e}"}),
                                        width="stretch", hide_index=True,
                                    )
                                else:
                                    # Backward-compatible display for older output files
                                    dt_df = pd.DataFrame(
                                        [r[:4] for r in dt_rows],
                                        columns=["Distance (m)", "TEDE (rem)",
                                                 "Limit (rem)", "Margin (× limit)"])
                                    dt_df = _clean_dose_display_df(dt_df)
                                    dt_df["Distance (m)"] = dt_df["Distance (m)"].astype("Int64")
                                    st.dataframe(
                                        dt_df.style.format(
                                            {"TEDE (rem)": "{:.4e}",
                                             "Margin (× limit)": "{:,.1f}"}),
                                        width="stretch", hide_index=True,
                                    )

                except Exception as _de:
                    st.warning(f"Could not parse dose results: {_de}")

        # ── Iodine Spike Dose (always shown when available) ──────────────────
        isp_rows = load_iodine_spike(run_case, run_dir=_run_dir)
        if isp_rows:
            with st.expander("💧  Iodine Spike — Pre-existing Coolant Activity",
                             expanded=True):
                try:
                    def _isp_val(key):
                        row = next((r for r in isp_rows if r[0]==key), None)
                        return row[1] if row and len(row)>1 else None

                    act   = _isp_val("Coolant activity (uCi/g)")
                    mass  = _isp_val("Primary mass (kg)")
                    sp_ci = _isp_val("Spike inventory (Ci)")
                    mult  = _isp_val("Spike multiplier")
                    model = _isp_val("Model")
                    scram = _isp_val("Scram occurred")
                    porv  = _isp_val("PORV opened")
                    dnb   = _isp_val("DNB occurred")
                    n_dnb = _isp_val("N rods DNB/dryout")
                    f_eq  = _isp_val("Equilibrium spike frac")
                    f_acc = _isp_val("Accident spike frac (DNB)")

                    # Input metrics row
                    c1, c2, c3 = st.columns(3)
                    c1.metric("Coolant Activity",
                              f"{act:.3f} μCi/g" if act is not None else "—")
                    c2.metric("Primary Mass",
                              f"{mass:,.0f} kg" if mass is not None else "—")
                    c3.metric("Spike Inventory",
                              f"{sp_ci:.3e} Ci" if sp_ci is not None else "—",
                              f"{mult}× coolant activity")

                    # Trigger flags
                    st.markdown(
                        f"**Model:** {model or '—'} &nbsp;|&nbsp; "
                        f"{'✅' if scram=='True' else '❌'} Scram &nbsp;"
                        f"{'✅' if porv=='True' else '❌'} PORV &nbsp;"
                        f"{'✅' if dnb=='True' else '❌'} DNB"
                        + (f" ({int(float(n_dnb)):,} rods)"
                           if n_dnb and float(n_dnb)>0 else ""),
                        unsafe_allow_html=True,
                    )

                    # Summary TEDE table
                    summary = [r[:5] for r in isp_rows
                               if r[0] in ("EAB","LPZ","Control Room")]
                    if summary:
                        sum_df = pd.DataFrame(summary,
                            columns=["Location","TEDE (rem)","Limit (rem)",
                                     "Margin (rem)","Result"])
                        sum_df = _clean_dose_display_df(sum_df)
                        def _cr(v):
                            return ("color: green; font-weight: bold"
                                    if v=="PASS" else "color: red; font-weight: bold")
                        st.dataframe(
                            sum_df.style.map(_cr, subset=["Result"])
                                  .format({"TEDE (rem)":"{:.4e}",
                                           "Margin (rem)":"{:.4e}"}),
                            width="stretch", hide_index=True)

                    # Group contributions
                    grp_hdr = next((i for i,r in enumerate(isp_rows)
                                    if r[0]=="Group"), None)
                    if grp_hdr is not None:
                        grp_rows = []
                        for r in isp_rows[grp_hdr+1:]:
                            if not any(r): break
                            grp_rows.append(r)
                        if grp_rows:
                            grp_df = pd.DataFrame(
                                [r[:6] for r in grp_rows],
                                columns=["Group","Released (Ci)","EAB (rem)",
                                         "LPZ (rem)","CR (rem)","Release frac"])
                            grp_df = _clean_dose_display_df(grp_df)
                            st.markdown("**Group contributions:**")
                            st.dataframe(
                                grp_df.style.format(
                                    {"Released (Ci)":"{:.4e}","EAB (rem)":"{:.4e}",
                                     "LPZ (rem)":"{:.4e}","CR (rem)":"{:.4e}",
                                     "Release frac":"{:.4e}"}),
                                width="stretch", hide_index=True)

                    # chi/Q caption
                    xq_eab = _isp_val("chi/Q EAB (s/m3)")
                    rt     = _isp_val("Runtime (ms)")
                    meta = []
                    if xq_eab: meta.append(f"χ/Q EAB = {xq_eab:.3e} s/m³")
                    if rt:     meta.append(f"runtime = {rt:.1f} ms")
                    if meta:   st.caption("  ·  ".join(meta))

                    # Distance table
                    fit_row = next((r for r in isp_rows if r[0]=="Power-law fit"),None)
                    if fit_row and fit_row[1]:
                        st.markdown("---")
                        st.markdown("**TEDE vs Distance:**")
                        st.code(fit_row[1], language=None)
                        dist_hdr = next((i for i,r in enumerate(isp_rows)
                                         if r[0]=="Distance (m)"),None)
                        if dist_hdr is not None:
                            dt_rows = []
                            for r in isp_rows[dist_hdr+1:]:
                                if not any(r): break
                                dt_rows.append(r)
                            if dt_rows:
                                dt_df = pd.DataFrame(
                                    [r[:4] for r in dt_rows],
                                    columns=["Distance (m)","TEDE (rem)",
                                             "Limit (rem)","Margin (× limit)"])
                                dt_df = _clean_dose_display_df(dt_df)
                                dt_df["Distance (m)"] = dt_df["Distance (m)"].astype("Int64")
                                st.dataframe(
                                    dt_df.style.format(
                                        {"TEDE (rem)":"{:.4e}",
                                         "Margin (× limit)":"{:,.1f}"}),
                                    width="stretch", hide_index=True)
                except Exception as _isp_e:
                    st.warning(f"Could not parse iodine spike results: {_isp_e}")

        st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)



        # Downloads
        st.markdown('<div class="hdiv"></div>', unsafe_allow_html=True)
        dl_col1, dl_col2 = st.columns(2)
        with dl_col1:
            # Serve the CSV that was already written to disk by post_processing
            csv_name = f"{run_case}_out.csv"
            csv_path = (_run_dir / csv_name if _run_dir else None)
            if csv_path is None or not csv_path.exists():
                csv_path = WORK_DIR / csv_name
            if csv_path.exists():
                with open(csv_path, "rb") as _csv:
                    st.download_button(
                        "⬇  Download time-series CSV",
                        data=_csv.read(),
                        file_name=csv_name,
                        mime="text/csv",
                        width="stretch",
                    )
            else:
                st.button("⬇  Download time-series CSV",
                          disabled=True, width="stretch",
                          help="Run the simulation to generate the CSV")
        with dl_col2:
            # Look in run_dir first, then fall back to WORK_DIR
            pdf_name = f"{run_case}_figures.pdf"
            pdf_path = (_run_dir / pdf_name if _run_dir else None)
            if pdf_path is None or not pdf_path.exists():
                pdf_path = WORK_DIR / pdf_name
            if pdf_path.exists():
                with open(pdf_path, "rb") as _pdf:
                    st.download_button(
                        "⬇  Download figures PDF",
                        data=_pdf.read(),
                        file_name=pdf_name,
                        mime="application/pdf",
                        width="stretch",
                    )
            else:
                st.button("⬇  Download figures PDF",
                          disabled=True, width="stretch",
                          help="Run the simulation to generate the PDF")

